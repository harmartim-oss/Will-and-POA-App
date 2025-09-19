"""
Enhanced Power of Attorney Document Generator
Specialized service for generating Power of Attorney for Property and Personal Care documents
compliant with Ontario's Substitute Decisions Act
"""

import os
import json
import logging
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime, date
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .nlp_service import get_nlp_service

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class POAParty:
    """Represents a party in a Power of Attorney document"""
    full_name: str
    address: str
    city: str
    province: str
    postal_code: str
    phone: Optional[str] = None
    email: Optional[str] = None
    relationship: Optional[str] = None
    date_of_birth: Optional[str] = None

@dataclass
class POAWitness:
    """Represents a witness for POA execution"""
    full_name: str
    address: str
    city: str
    province: str
    postal_code: str
    occupation: Optional[str] = None
    phone: Optional[str] = None

@dataclass
class POARestrictions:
    """Represents restrictions and limitations on POA powers"""
    general_restrictions: List[str]
    specific_limitations: List[str]
    prohibited_actions: List[str]
    conditions: List[str]
    effective_date: Optional[str] = None
    termination_conditions: List[str] = None

@dataclass
class POAPropertyPowers:
    """Specific powers for Property POA"""
    real_estate_powers: List[str]
    financial_powers: List[str]
    business_powers: List[str]
    investment_powers: List[str]
    banking_powers: List[str]
    tax_powers: List[str]
    insurance_powers: List[str]
    legal_powers: List[str]

@dataclass
class POACarePowers:
    """Specific powers for Personal Care POA"""
    healthcare_powers: List[str]
    residential_powers: List[str]
    personal_care_powers: List[str]
    social_powers: List[str]
    communication_powers: List[str]
    safety_powers: List[str]

@dataclass
class POADocument:
    """Complete Power of Attorney document structure"""
    document_type: str  # "property" or "personal_care"
    grantor: POAParty
    attorneys: List[POAParty]
    substitute_attorneys: List[POAParty]
    witnesses: List[POAWitness]
    is_continuing: bool
    restrictions: POARestrictions
    property_powers: Optional[POAPropertyPowers] = None
    care_powers: Optional[POACarePowers] = None
    special_instructions: List[str] = None
    compensation_clause: Optional[str] = None
    accounting_requirements: List[str] = None
    execution_date: Optional[str] = None
    notarization_required: bool = False

class PowerOfAttorneyGenerator:
    """
    Enhanced generator for Power of Attorney documents compliant with Ontario law
    """
    
    def __init__(self):
        self.nlp_service = get_nlp_service()
        
        # Ontario legal requirements
        self.ontario_requirements = {
            'property_poa': {
                'minimum_age': 18,
                'witness_requirements': {
                    'minimum_witnesses': 2,
                    'restrictions': [
                        'Cannot be the attorney named in the document',
                        'Cannot be the spouse or partner of the attorney',
                        'Cannot be a child of the grantor if the attorney is the spouse or partner of the grantor',
                        'Must be at least 18 years old',
                        'Must be mentally capable'
                    ]
                },
                'execution_requirements': [
                    'Must be signed by grantor in presence of witnesses',
                    'Witnesses must sign in presence of grantor and each other',
                    'All signatures must be on the same document'
                ]
            },
            'personal_care_poa': {
                'minimum_age': 16,
                'witness_requirements': {
                    'minimum_witnesses': 2,
                    'restrictions': [
                        'Cannot be the attorney named in the document',
                        'Cannot be the spouse or partner of the attorney',
                        'Must be at least 18 years old',
                        'Must be mentally capable'
                    ]
                },
                'execution_requirements': [
                    'Must be signed by grantor in presence of witnesses',
                    'Witnesses must sign in presence of grantor and each other',
                    'All signatures must be on the same document'
                ]
            }
        }
        
        # Standard powers for different POA types
        self.standard_property_powers = {
            'real_estate_powers': [
                'Buy, sell, lease, mortgage, or otherwise deal with real estate',
                'Manage rental properties and collect rents',
                'Pay property taxes and maintenance costs',
                'Enter into agreements for property management',
                'Apply for building permits and approvals'
            ],
            'financial_powers': [
                'Operate bank accounts and investment accounts',
                'Make deposits and withdrawals',
                'Transfer funds between accounts',
                'Apply for loans and credit facilities',
                'Manage credit cards and lines of credit'
            ],
            'business_powers': [
                'Operate and manage business interests',
                'Enter into contracts on behalf of the grantor',
                'Hire and dismiss employees',
                'Make business decisions and investments',
                'File business tax returns and regulatory filings'
            ],
            'investment_powers': [
                'Buy, sell, and manage investments',
                'Make investment decisions',
                'Receive dividends and interest',
                'Exercise voting rights in corporations',
                'Manage retirement accounts and pensions'
            ],
            'banking_powers': [
                'Open and close bank accounts',
                'Endorse and deposit cheques',
                'Access safety deposit boxes',
                'Arrange for direct deposits and automatic payments',
                'Obtain bank statements and financial records'
            ],
            'tax_powers': [
                'File income tax returns',
                'Claim deductions and credits',
                'Correspond with tax authorities',
                'Pay taxes and penalties',
                'Represent the grantor in tax matters'
            ],
            'insurance_powers': [
                'Purchase, maintain, and cancel insurance policies',
                'File insurance claims',
                'Receive insurance proceeds',
                'Change beneficiaries on insurance policies',
                'Negotiate insurance settlements'
            ],
            'legal_powers': [
                'Retain and instruct lawyers',
                'Commence and defend legal proceedings',
                'Settle legal disputes',
                'Execute legal documents',
                'Appear in court on behalf of the grantor'
            ]
        }
        
        self.standard_care_powers = {
            'healthcare_powers': [
                'Make decisions about medical and dental treatment',
                'Consent to or refuse medical procedures',
                'Choose healthcare providers and facilities',
                'Access medical records and information',
                'Make end-of-life care decisions'
            ],
            'residential_powers': [
                'Choose where the grantor will live',
                'Arrange for home care services',
                'Make decisions about nursing home placement',
                'Manage household affairs',
                'Arrange for home maintenance and repairs'
            ],
            'personal_care_powers': [
                'Make decisions about nutrition and diet',
                'Arrange for personal hygiene and grooming',
                'Choose clothing and personal items',
                'Make decisions about daily activities',
                'Arrange for transportation'
            ],
            'social_powers': [
                'Make decisions about social activities',
                'Arrange for visits with family and friends',
                'Make decisions about religious observances',
                'Choose recreational activities',
                'Manage social relationships'
            ],
            'communication_powers': [
                'Make decisions about communication methods',
                'Access mail and correspondence',
                'Make decisions about telephone and internet use',
                'Communicate with healthcare providers',
                'Make decisions about information sharing'
            ],
            'safety_powers': [
                'Make decisions about personal safety',
                'Arrange for security measures',
                'Make decisions about driving and transportation',
                'Ensure safe living conditions',
                'Take protective measures when necessary'
            ]
        }
    
    def generate_poa_document(self, poa_data: Dict[str, Any]) -> POADocument:
        """
        Generate a complete POA document from user input
        
        Args:
            poa_data: Dictionary containing POA information
            
        Returns:
            POADocument object with complete document structure
        """
        try:
            # Parse basic information
            document_type = poa_data.get('document_type', 'property')
            
            # Create grantor
            grantor_data = poa_data.get('grantor', {})
            grantor = POAParty(**grantor_data)
            
            # Create attorneys
            attorneys = []
            for attorney_data in poa_data.get('attorneys', []):
                attorneys.append(POAParty(**attorney_data))
            
            # Create substitute attorneys
            substitute_attorneys = []
            for sub_attorney_data in poa_data.get('substitute_attorneys', []):
                substitute_attorneys.append(POAParty(**sub_attorney_data))
            
            # Create witnesses
            witnesses = []
            for witness_data in poa_data.get('witnesses', []):
                witnesses.append(POAWitness(**witness_data))
            
            # Create restrictions
            restrictions_data = poa_data.get('restrictions', {})
            restrictions = POARestrictions(
                general_restrictions=restrictions_data.get('general_restrictions', []),
                specific_limitations=restrictions_data.get('specific_limitations', []),
                prohibited_actions=restrictions_data.get('prohibited_actions', []),
                conditions=restrictions_data.get('conditions', []),
                effective_date=restrictions_data.get('effective_date'),
                termination_conditions=restrictions_data.get('termination_conditions', [])
            )
            
            # Create powers based on document type
            property_powers = None
            care_powers = None
            
            if document_type == 'property':
                powers_data = poa_data.get('property_powers', {})
                property_powers = POAPropertyPowers(
                    real_estate_powers=powers_data.get('real_estate_powers', self.standard_property_powers['real_estate_powers']),
                    financial_powers=powers_data.get('financial_powers', self.standard_property_powers['financial_powers']),
                    business_powers=powers_data.get('business_powers', self.standard_property_powers['business_powers']),
                    investment_powers=powers_data.get('investment_powers', self.standard_property_powers['investment_powers']),
                    banking_powers=powers_data.get('banking_powers', self.standard_property_powers['banking_powers']),
                    tax_powers=powers_data.get('tax_powers', self.standard_property_powers['tax_powers']),
                    insurance_powers=powers_data.get('insurance_powers', self.standard_property_powers['insurance_powers']),
                    legal_powers=powers_data.get('legal_powers', self.standard_property_powers['legal_powers'])
                )
            
            elif document_type == 'personal_care':
                powers_data = poa_data.get('care_powers', {})
                care_powers = POACarePowers(
                    healthcare_powers=powers_data.get('healthcare_powers', self.standard_care_powers['healthcare_powers']),
                    residential_powers=powers_data.get('residential_powers', self.standard_care_powers['residential_powers']),
                    personal_care_powers=powers_data.get('personal_care_powers', self.standard_care_powers['personal_care_powers']),
                    social_powers=powers_data.get('social_powers', self.standard_care_powers['social_powers']),
                    communication_powers=powers_data.get('communication_powers', self.standard_care_powers['communication_powers']),
                    safety_powers=powers_data.get('safety_powers', self.standard_care_powers['safety_powers'])
                )
            
            # Create complete document
            poa_document = POADocument(
                document_type=document_type,
                grantor=grantor,
                attorneys=attorneys,
                substitute_attorneys=substitute_attorneys,
                witnesses=witnesses,
                is_continuing=poa_data.get('is_continuing', True),
                restrictions=restrictions,
                property_powers=property_powers,
                care_powers=care_powers,
                special_instructions=poa_data.get('special_instructions', []),
                compensation_clause=poa_data.get('compensation_clause'),
                accounting_requirements=poa_data.get('accounting_requirements', []),
                execution_date=poa_data.get('execution_date'),
                notarization_required=poa_data.get('notarization_required', False)
            )
            
            return poa_document
            
        except Exception as e:
            logger.error(f"Error generating POA document: {e}")
            raise
    
    def validate_poa_document(self, poa_document: POADocument) -> Dict[str, Any]:
        """
        Validate POA document against Ontario legal requirements
        
        Args:
            poa_document: POADocument to validate
            
        Returns:
            Dictionary with validation results
        """
        validation_results = {
            'is_valid': True,
            'errors': [],
            'warnings': [],
            'requirements_met': [],
            'suggestions': []
        }
        
        try:
            # Get requirements for document type
            if poa_document.document_type == 'property':
                requirements = self.ontario_requirements['property_poa']
            else:
                requirements = self.ontario_requirements['personal_care_poa']
            
            # Validate grantor age
            if poa_document.grantor.date_of_birth:
                try:
                    birth_date = datetime.strptime(poa_document.grantor.date_of_birth, '%Y-%m-%d').date()
                    age = (date.today() - birth_date).days // 365
                    
                    if age < requirements['minimum_age']:
                        validation_results['errors'].append(
                            f"Grantor must be at least {requirements['minimum_age']} years old for {poa_document.document_type} POA"
                        )
                        validation_results['is_valid'] = False
                    else:
                        validation_results['requirements_met'].append("Age requirement satisfied")
                except ValueError:
                    validation_results['warnings'].append("Invalid date of birth format")
            else:
                validation_results['warnings'].append("Date of birth not provided - cannot verify age requirement")
            
            # Validate attorneys
            if not poa_document.attorneys:
                validation_results['errors'].append("At least one attorney must be appointed")
                validation_results['is_valid'] = False
            else:
                validation_results['requirements_met'].append("Attorney(s) appointed")
            
            # Validate witnesses
            witness_count = len(poa_document.witnesses)
            required_witnesses = requirements['witness_requirements']['minimum_witnesses']
            
            if witness_count < required_witnesses:
                validation_results['errors'].append(
                    f"Minimum {required_witnesses} witnesses required, only {witness_count} provided"
                )
                validation_results['is_valid'] = False
            else:
                validation_results['requirements_met'].append(f"Sufficient witnesses ({witness_count})")
            
            # Validate witness restrictions
            for i, witness in enumerate(poa_document.witnesses):
                # Check if witness is an attorney
                for attorney in poa_document.attorneys:
                    if witness.full_name.lower() == attorney.full_name.lower():
                        validation_results['errors'].append(
                            f"Witness {i+1} ({witness.full_name}) cannot be an attorney"
                        )
                        validation_results['is_valid'] = False
            
            # Validate powers
            if poa_document.document_type == 'property' and not poa_document.property_powers:
                validation_results['warnings'].append("No specific property powers defined - using standard powers")
            elif poa_document.document_type == 'personal_care' and not poa_document.care_powers:
                validation_results['warnings'].append("No specific care powers defined - using standard powers")
            
            # Validate continuing POA requirements
            if poa_document.is_continuing:
                validation_results['requirements_met'].append("Continuing POA provisions included")
                if poa_document.document_type == 'personal_care':
                    validation_results['suggestions'].append(
                        "Consider including specific instructions for incapacity determination"
                    )
            
            # Validate execution requirements
            if not poa_document.execution_date:
                validation_results['warnings'].append("Execution date not specified")
            
            # Additional suggestions
            if not poa_document.substitute_attorneys:
                validation_results['suggestions'].append("Consider appointing substitute attorneys")
            
            if not poa_document.special_instructions:
                validation_results['suggestions'].append("Consider adding special instructions or preferences")
            
            if poa_document.document_type == 'property' and not poa_document.accounting_requirements:
                validation_results['suggestions'].append("Consider specifying accounting requirements for attorneys")
            
        except Exception as e:
            logger.error(f"Error validating POA document: {e}")
            validation_results['errors'].append(f"Validation error: {str(e)}")
            validation_results['is_valid'] = False
        
        return validation_results
    
    def generate_pdf(self, poa_document: POADocument, output_path: str) -> str:
        """
        Generate PDF version of POA document
        
        Args:
            poa_document: POADocument to convert
            output_path: Path for output PDF file
            
        Returns:
            Path to generated PDF file
        """
        try:
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=letter,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=12,
                spaceBefore=12,
                fontName='Helvetica-Bold'
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6,
                alignment=TA_JUSTIFY,
                fontName='Helvetica'
            )
            
            # Build document content
            story = []
            
            # Title
            if poa_document.document_type == 'property':
                title = "CONTINUING POWER OF ATTORNEY FOR PROPERTY" if poa_document.is_continuing else "POWER OF ATTORNEY FOR PROPERTY"
            else:
                title = "POWER OF ATTORNEY FOR PERSONAL CARE"
            
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))
            
            # Grantor information
            story.append(Paragraph("GRANTOR INFORMATION", heading_style))
            grantor_text = f"""
            I, <b>{poa_document.grantor.full_name}</b>, of {poa_document.grantor.address}, 
            {poa_document.grantor.city}, {poa_document.grantor.province} {poa_document.grantor.postal_code}, 
            being of sound mind and legal age, hereby make this Power of Attorney.
            """
            story.append(Paragraph(grantor_text, body_style))
            story.append(Spacer(1, 12))
            
            # Attorney appointment
            story.append(Paragraph("APPOINTMENT OF ATTORNEY(S)", heading_style))
            
            for i, attorney in enumerate(poa_document.attorneys, 1):
                attorney_text = f"""
                {i}. I hereby appoint <b>{attorney.full_name}</b>, of {attorney.address}, 
                {attorney.city}, {attorney.province} {attorney.postal_code}, 
                to be my attorney for {poa_document.document_type.replace('_', ' ')}.
                """
                story.append(Paragraph(attorney_text, body_style))
            
            story.append(Spacer(1, 12))
            
            # Substitute attorneys
            if poa_document.substitute_attorneys:
                story.append(Paragraph("SUBSTITUTE ATTORNEY(S)", heading_style))
                for i, sub_attorney in enumerate(poa_document.substitute_attorneys, 1):
                    sub_text = f"""
                    {i}. If my attorney is unable or unwilling to act, I appoint <b>{sub_attorney.full_name}</b>, 
                    of {sub_attorney.address}, {sub_attorney.city}, {sub_attorney.province} {sub_attorney.postal_code}, 
                    to be my substitute attorney.
                    """
                    story.append(Paragraph(sub_text, body_style))
                story.append(Spacer(1, 12))
            
            # Powers granted
            story.append(Paragraph("POWERS GRANTED", heading_style))
            
            if poa_document.document_type == 'property' and poa_document.property_powers:
                self._add_property_powers_to_pdf(story, poa_document.property_powers, body_style)
            elif poa_document.document_type == 'personal_care' and poa_document.care_powers:
                self._add_care_powers_to_pdf(story, poa_document.care_powers, body_style)
            
            # Restrictions and limitations
            if poa_document.restrictions.general_restrictions or poa_document.restrictions.specific_limitations:
                story.append(Paragraph("RESTRICTIONS AND LIMITATIONS", heading_style))
                
                if poa_document.restrictions.general_restrictions:
                    story.append(Paragraph("General Restrictions:", body_style))
                    for restriction in poa_document.restrictions.general_restrictions:
                        story.append(Paragraph(f"• {restriction}", body_style))
                
                if poa_document.restrictions.specific_limitations:
                    story.append(Paragraph("Specific Limitations:", body_style))
                    for limitation in poa_document.restrictions.specific_limitations:
                        story.append(Paragraph(f"• {limitation}", body_style))
                
                story.append(Spacer(1, 12))
            
            # Special instructions
            if poa_document.special_instructions:
                story.append(Paragraph("SPECIAL INSTRUCTIONS", heading_style))
                for instruction in poa_document.special_instructions:
                    story.append(Paragraph(f"• {instruction}", body_style))
                story.append(Spacer(1, 12))
            
            # Continuing POA clause
            if poa_document.is_continuing:
                story.append(Paragraph("CONTINUING POWER OF ATTORNEY", heading_style))
                continuing_text = """
                This Power of Attorney shall continue to be effective if I become mentally incapable 
                of managing my property, in accordance with the Substitute Decisions Act, 1992.
                """
                story.append(Paragraph(continuing_text, body_style))
                story.append(Spacer(1, 12))
            
            # Execution section
            story.append(Paragraph("EXECUTION", heading_style))
            
            execution_text = f"""
            IN WITNESS WHEREOF, I have executed this Power of Attorney on this _____ day of 
            _____________, 20____.
            """
            story.append(Paragraph(execution_text, body_style))
            story.append(Spacer(1, 24))
            
            # Signature lines
            signature_data = [
                ['', ''],
                ['_' * 40, '_' * 40],
                [f'{poa_document.grantor.full_name}', 'Date'],
                ['(Grantor)', '']
            ]
            
            signature_table = Table(signature_data, colWidths=[3*inch, 2*inch])
            signature_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 2), (-1, 2), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
            ]))
            
            story.append(signature_table)
            story.append(Spacer(1, 24))
            
            # Witness section
            story.append(Paragraph("WITNESSES", heading_style))
            
            witness_text = """
            Signed by the above-named Grantor in our presence, and by us in the presence 
            of the Grantor and each other:
            """
            story.append(Paragraph(witness_text, body_style))
            story.append(Spacer(1, 12))
            
            # Witness signature table
            for i, witness in enumerate(poa_document.witnesses, 1):
                witness_data = [
                    ['', ''],
                    ['_' * 40, '_' * 40],
                    [f'Witness {i}: {witness.full_name}', 'Date'],
                    [f'Address: {witness.address}', ''],
                    [f'{witness.city}, {witness.province} {witness.postal_code}', '']
                ]
                
                witness_table = Table(witness_data, colWidths=[3*inch, 2*inch])
                witness_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 2), (-1, 2), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                ]))
                
                story.append(witness_table)
                story.append(Spacer(1, 12))
            
            # Legal notice
            story.append(PageBreak())
            story.append(Paragraph("IMPORTANT LEGAL INFORMATION", heading_style))
            
            legal_notice = """
            This Power of Attorney has been prepared in accordance with the laws of Ontario, Canada. 
            It is recommended that you seek independent legal advice before executing this document. 
            
            The Substitute Decisions Act, 1992 governs Powers of Attorney in Ontario. This document 
            should be executed in the presence of two witnesses who meet the requirements set out 
            in the Act.
            
            For more information about Powers of Attorney in Ontario, contact the Office of the 
            Public Guardian and Trustee or consult with a qualified legal professional.
            """
            story.append(Paragraph(legal_notice, body_style))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            raise
    
    def _add_property_powers_to_pdf(self, story: List, property_powers: POAPropertyPowers, style):
        """Add property powers to PDF story"""
        powers_sections = [
            ("Real Estate Powers", property_powers.real_estate_powers),
            ("Financial Powers", property_powers.financial_powers),
            ("Business Powers", property_powers.business_powers),
            ("Investment Powers", property_powers.investment_powers),
            ("Banking Powers", property_powers.banking_powers),
            ("Tax Powers", property_powers.tax_powers),
            ("Insurance Powers", property_powers.insurance_powers),
            ("Legal Powers", property_powers.legal_powers)
        ]
        
        for section_title, powers in powers_sections:
            if powers:
                story.append(Paragraph(f"<b>{section_title}:</b>", style))
                for power in powers:
                    story.append(Paragraph(f"• {power}", style))
                story.append(Spacer(1, 6))
    
    def _add_care_powers_to_pdf(self, story: List, care_powers: POACarePowers, style):
        """Add personal care powers to PDF story"""
        powers_sections = [
            ("Healthcare Powers", care_powers.healthcare_powers),
            ("Residential Powers", care_powers.residential_powers),
            ("Personal Care Powers", care_powers.personal_care_powers),
            ("Social Powers", care_powers.social_powers),
            ("Communication Powers", care_powers.communication_powers),
            ("Safety Powers", care_powers.safety_powers)
        ]
        
        for section_title, powers in powers_sections:
            if powers:
                story.append(Paragraph(f"<b>{section_title}:</b>", style))
                for power in powers:
                    story.append(Paragraph(f"• {power}", style))
                story.append(Spacer(1, 6))
    
    def generate_word_document(self, poa_document: POADocument, output_path: str) -> str:
        """
        Generate Word document version of POA
        
        Args:
            poa_document: POADocument to convert
            output_path: Path for output Word file
            
        Returns:
            Path to generated Word file
        """
        try:
            # Create Word document
            doc = docx.Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Title
            if poa_document.document_type == 'property':
                title = "CONTINUING POWER OF ATTORNEY FOR PROPERTY" if poa_document.is_continuing else "POWER OF ATTORNEY FOR PROPERTY"
            else:
                title = "POWER OF ATTORNEY FOR PERSONAL CARE"
            
            title_paragraph = doc.add_paragraph(title)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.runs[0]
            title_run.bold = True
            title_run.font.size = docx.shared.Pt(16)
            
            doc.add_paragraph()  # Empty line
            
            # Grantor information
            doc.add_heading('GRANTOR INFORMATION', level=2)
            grantor_text = f"""I, {poa_document.grantor.full_name}, of {poa_document.grantor.address}, {poa_document.grantor.city}, {poa_document.grantor.province} {poa_document.grantor.postal_code}, being of sound mind and legal age, hereby make this Power of Attorney."""
            doc.add_paragraph(grantor_text)
            
            # Attorney appointment
            doc.add_heading('APPOINTMENT OF ATTORNEY(S)', level=2)
            for i, attorney in enumerate(poa_document.attorneys, 1):
                attorney_text = f"""{i}. I hereby appoint {attorney.full_name}, of {attorney.address}, {attorney.city}, {attorney.province} {attorney.postal_code}, to be my attorney for {poa_document.document_type.replace('_', ' ')}."""
                doc.add_paragraph(attorney_text)
            
            # Substitute attorneys
            if poa_document.substitute_attorneys:
                doc.add_heading('SUBSTITUTE ATTORNEY(S)', level=2)
                for i, sub_attorney in enumerate(poa_document.substitute_attorneys, 1):
                    sub_text = f"""{i}. If my attorney is unable or unwilling to act, I appoint {sub_attorney.full_name}, of {sub_attorney.address}, {sub_attorney.city}, {sub_attorney.province} {sub_attorney.postal_code}, to be my substitute attorney."""
                    doc.add_paragraph(sub_text)
            
            # Powers granted
            doc.add_heading('POWERS GRANTED', level=2)
            
            if poa_document.document_type == 'property' and poa_document.property_powers:
                self._add_property_powers_to_word(doc, poa_document.property_powers)
            elif poa_document.document_type == 'personal_care' and poa_document.care_powers:
                self._add_care_powers_to_word(doc, poa_document.care_powers)
            
            # Restrictions and limitations
            if poa_document.restrictions.general_restrictions or poa_document.restrictions.specific_limitations:
                doc.add_heading('RESTRICTIONS AND LIMITATIONS', level=2)
                
                if poa_document.restrictions.general_restrictions:
                    doc.add_paragraph('General Restrictions:').runs[0].bold = True
                    for restriction in poa_document.restrictions.general_restrictions:
                        doc.add_paragraph(f'• {restriction}')
                
                if poa_document.restrictions.specific_limitations:
                    doc.add_paragraph('Specific Limitations:').runs[0].bold = True
                    for limitation in poa_document.restrictions.specific_limitations:
                        doc.add_paragraph(f'• {limitation}')
            
            # Special instructions
            if poa_document.special_instructions:
                doc.add_heading('SPECIAL INSTRUCTIONS', level=2)
                for instruction in poa_document.special_instructions:
                    doc.add_paragraph(f'• {instruction}')
            
            # Continuing POA clause
            if poa_document.is_continuing:
                doc.add_heading('CONTINUING POWER OF ATTORNEY', level=2)
                continuing_text = """This Power of Attorney shall continue to be effective if I become mentally incapable of managing my property, in accordance with the Substitute Decisions Act, 1992."""
                doc.add_paragraph(continuing_text)
            
            # Execution section
            doc.add_heading('EXECUTION', level=2)
            execution_text = """IN WITNESS WHEREOF, I have executed this Power of Attorney on this _____ day of _____________, 20____."""
            doc.add_paragraph(execution_text)
            
            doc.add_paragraph()  # Empty line
            
            # Signature section
            signature_table = doc.add_table(rows=3, cols=2)
            signature_table.style = 'Table Grid'
            
            signature_table.cell(0, 0).text = ''
            signature_table.cell(0, 1).text = ''
            signature_table.cell(1, 0).text = '_' * 40
            signature_table.cell(1, 1).text = '_' * 40
            signature_table.cell(2, 0).text = f'{poa_document.grantor.full_name} (Grantor)'
            signature_table.cell(2, 1).text = 'Date'
            
            # Witness section
            doc.add_heading('WITNESSES', level=2)
            witness_text = """Signed by the above-named Grantor in our presence, and by us in the presence of the Grantor and each other:"""
            doc.add_paragraph(witness_text)
            
            for i, witness in enumerate(poa_document.witnesses, 1):
                doc.add_paragraph()  # Empty line
                witness_table = doc.add_table(rows=4, cols=2)
                witness_table.style = 'Table Grid'
                
                witness_table.cell(0, 0).text = ''
                witness_table.cell(0, 1).text = ''
                witness_table.cell(1, 0).text = '_' * 40
                witness_table.cell(1, 1).text = '_' * 40
                witness_table.cell(2, 0).text = f'Witness {i}: {witness.full_name}'
                witness_table.cell(2, 1).text = 'Date'
                witness_table.cell(3, 0).text = f'Address: {witness.address}, {witness.city}, {witness.province} {witness.postal_code}'
                witness_table.cell(3, 1).text = ''
            
            # Legal notice
            doc.add_page_break()
            doc.add_heading('IMPORTANT LEGAL INFORMATION', level=2)
            
            legal_notice = """This Power of Attorney has been prepared in accordance with the laws of Ontario, Canada. It is recommended that you seek independent legal advice before executing this document.

The Substitute Decisions Act, 1992 governs Powers of Attorney in Ontario. This document should be executed in the presence of two witnesses who meet the requirements set out in the Act.

For more information about Powers of Attorney in Ontario, contact the Office of the Public Guardian and Trustee or consult with a qualified legal professional."""
            
            doc.add_paragraph(legal_notice)
            
            # Save document
            doc.save(output_path)
            
            logger.info(f"Word document generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            raise
    
    def _add_property_powers_to_word(self, doc, property_powers: POAPropertyPowers):
        """Add property powers to Word document"""
        powers_sections = [
            ("Real Estate Powers", property_powers.real_estate_powers),
            ("Financial Powers", property_powers.financial_powers),
            ("Business Powers", property_powers.business_powers),
            ("Investment Powers", property_powers.investment_powers),
            ("Banking Powers", property_powers.banking_powers),
            ("Tax Powers", property_powers.tax_powers),
            ("Insurance Powers", property_powers.insurance_powers),
            ("Legal Powers", property_powers.legal_powers)
        ]
        
        for section_title, powers in powers_sections:
            if powers:
                paragraph = doc.add_paragraph(f'{section_title}:')
                paragraph.runs[0].bold = True
                for power in powers:
                    doc.add_paragraph(f'• {power}')
    
    def _add_care_powers_to_word(self, doc, care_powers: POACarePowers):
        """Add personal care powers to Word document"""
        powers_sections = [
            ("Healthcare Powers", care_powers.healthcare_powers),
            ("Residential Powers", care_powers.residential_powers),
            ("Personal Care Powers", care_powers.personal_care_powers),
            ("Social Powers", care_powers.social_powers),
            ("Communication Powers", care_powers.communication_powers),
            ("Safety Powers", care_powers.safety_powers)
        ]
        
        for section_title, powers in powers_sections:
            if powers:
                paragraph = doc.add_paragraph(f'{section_title}:')
                paragraph.runs[0].bold = True
                for power in powers:
                    doc.add_paragraph(f'• {power}')
    
    def analyze_poa_with_nlp(self, poa_document: POADocument) -> Dict[str, Any]:
        """
        Analyze POA document using NLP for improvements and compliance
        
        Args:
            poa_document: POADocument to analyze
            
        Returns:
            Dictionary with NLP analysis results
        """
        try:
            # Convert document to text for analysis
            document_text = self._poa_to_text(poa_document)
            
            # Perform NLP analysis
            analysis = self.nlp_service.analyze_legal_text(document_text)
            
            # Generate POA-specific suggestions
            poa_suggestions = self._generate_poa_suggestions(poa_document, analysis)
            
            return {
                'nlp_analysis': {
                    'entities': [
                        {
                            'text': entity.text,
                            'label': entity.label,
                            'description': entity.description
                        }
                        for entity in analysis.entities
                    ],
                    'sentiment': analysis.sentiment,
                    'readability_score': analysis.readability_score,
                    'complexity_score': analysis.complexity_score,
                    'legal_concepts': analysis.legal_concepts
                },
                'poa_specific_suggestions': poa_suggestions,
                'compliance_analysis': analysis.compliance_issues,
                'risk_factors': analysis.risk_factors,
                'document_metrics': {
                    'word_count': analysis.word_count,
                    'sentence_count': analysis.sentence_count
                }
            }
            
        except Exception as e:
            logger.error(f"Error analyzing POA with NLP: {e}")
            return {'error': str(e)}
    
    def _poa_to_text(self, poa_document: POADocument) -> str:
        """Convert POA document to text for NLP analysis"""
        text_parts = []
        
        # Add document type
        text_parts.append(f"Power of Attorney for {poa_document.document_type.replace('_', ' ')}")
        
        # Add grantor information
        text_parts.append(f"Grantor: {poa_document.grantor.full_name}")
        
        # Add attorneys
        for attorney in poa_document.attorneys:
            text_parts.append(f"Attorney: {attorney.full_name}")
        
        # Add powers
        if poa_document.property_powers:
            for powers_list in [
                poa_document.property_powers.real_estate_powers,
                poa_document.property_powers.financial_powers,
                poa_document.property_powers.business_powers,
                poa_document.property_powers.investment_powers,
                poa_document.property_powers.banking_powers,
                poa_document.property_powers.tax_powers,
                poa_document.property_powers.insurance_powers,
                poa_document.property_powers.legal_powers
            ]:
                text_parts.extend(powers_list)
        
        if poa_document.care_powers:
            for powers_list in [
                poa_document.care_powers.healthcare_powers,
                poa_document.care_powers.residential_powers,
                poa_document.care_powers.personal_care_powers,
                poa_document.care_powers.social_powers,
                poa_document.care_powers.communication_powers,
                poa_document.care_powers.safety_powers
            ]:
                text_parts.extend(powers_list)
        
        # Add restrictions and special instructions
        text_parts.extend(poa_document.restrictions.general_restrictions)
        text_parts.extend(poa_document.restrictions.specific_limitations)
        if poa_document.special_instructions:
            text_parts.extend(poa_document.special_instructions)
        
        return ' '.join(text_parts)
    
    def _generate_poa_suggestions(self, poa_document: POADocument, nlp_analysis) -> List[str]:
        """Generate POA-specific suggestions based on analysis"""
        suggestions = []
        
        # Check for missing elements
        if not poa_document.substitute_attorneys:
            suggestions.append("Consider appointing substitute attorneys in case primary attorneys cannot serve")
        
        if not poa_document.special_instructions:
            suggestions.append("Consider adding special instructions or preferences for your attorneys")
        
        if poa_document.document_type == 'property':
            if not poa_document.accounting_requirements:
                suggestions.append("Consider specifying accounting requirements for property management")
            
            if not poa_document.compensation_clause:
                suggestions.append("Consider addressing attorney compensation in the document")
        
        # Check for clarity issues
        if nlp_analysis.readability_score < 0.5:
            suggestions.append("Consider simplifying language for better clarity")
        
        if nlp_analysis.complexity_score > 0.8:
            suggestions.append("Document may be overly complex - consider breaking down complex clauses")
        
        # Check for completeness
        if len(poa_document.witnesses) < 2:
            suggestions.append("Ensure you have at least two qualified witnesses for execution")
        
        return suggestions

# Initialize global POA generator
poa_generator = PowerOfAttorneyGenerator()

def get_poa_generator() -> PowerOfAttorneyGenerator:
    """Get the global POA generator instance"""
    return poa_generator

