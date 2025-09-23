"""
AI-Enhanced Professional Document Generator
Generates Ontario-compliant legal documents with AI-powered enhancements
"""

import os
import tempfile
import json
from typing import Dict, List, Any, Optional
from dataclasses import asdict
from datetime import datetime
import logging

# Document generation imports (would use actual libraries in production)
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from enhanced_ai_legal_service import EnhancedAILegalService
from ontario_legal_kb import OntarioLegalKnowledgeBase

logger = logging.getLogger(__name__)

class AIEnhancedDocumentGenerator:
    """
    Professional document generator with AI enhancements and Ontario legal compliance
    """
    
    def __init__(self):
        self.ai_service = EnhancedAILegalService()
        self.legal_kb = OntarioLegalKnowledgeBase()
        self.style_config = self._load_style_configuration()
        
    def _load_style_configuration(self) -> Dict[str, Any]:
        """Load professional document styling configuration"""
        return {
            "fonts": {
                "main": "Times New Roman",
                "heading": "Arial",
                "legal": "Times New Roman"
            },
            "font_sizes": {
                "title": 16,
                "heading": 14,
                "body": 12,
                "footer": 10
            },
            "margins": {
                "top": 1.0,
                "bottom": 1.0,
                "left": 1.0,
                "right": 1.0
            },
            "line_spacing": 1.15,
            "colors": {
                "text": "000000",
                "heading": "1a1a1a",
                "accent": "2c5aa0"
            }
        }
    
    def generate_ontario_will(self, form_data: Dict[str, Any], 
                            enable_ai_enhancements: bool = True) -> Dict[str, Any]:
        """Generate AI-enhanced Ontario will document"""
        logger.info("Generating Ontario will with AI enhancements")
        
        # Analyze document with AI first
        ai_analysis = None
        if enable_ai_enhancements:
            ai_analysis = self.ai_service.analyze_document("will", form_data)
        
        # Create document
        if DOCX_AVAILABLE:
            doc = Document()
            self._apply_ontario_legal_styling(doc)
            
            # Generate content with AI enhancements
            content = self._generate_will_content(form_data, ai_analysis)
            
            # Add content to document
            self._add_will_sections(doc, content, form_data)
            
            # Save document
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            return {
                "document_path": temp_file.name,
                "document_type": "will",
                "format": "docx",
                "ai_analysis": asdict(ai_analysis) if ai_analysis else None,
                "generation_time": datetime.now().isoformat(),
                "compliance_score": ai_analysis.compliance_score if ai_analysis else None
            }
        else:
            # Fallback to text generation
            return self._generate_will_text_format(form_data, ai_analysis)
    
    def generate_ontario_poa(self, poa_type: str, form_data: Dict[str, Any],
                           enable_ai_enhancements: bool = True) -> Dict[str, Any]:
        """Generate AI-enhanced Ontario Power of Attorney document"""
        logger.info(f"Generating Ontario {poa_type} with AI enhancements")
        
        # Analyze document with AI
        ai_analysis = None
        if enable_ai_enhancements:
            ai_analysis = self.ai_service.analyze_document(poa_type, form_data)
        
        if DOCX_AVAILABLE:
            doc = Document()
            self._apply_ontario_legal_styling(doc)
            
            # Generate content
            content = self._generate_poa_content(poa_type, form_data, ai_analysis)
            
            # Add content to document
            self._add_poa_sections(doc, poa_type, content, form_data)
            
            # Save document
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            return {
                "document_path": temp_file.name,
                "document_type": poa_type,
                "format": "docx",
                "ai_analysis": asdict(ai_analysis) if ai_analysis else None,
                "generation_time": datetime.now().isoformat(),
                "compliance_score": ai_analysis.compliance_score if ai_analysis else None
            }
        else:
            return self._generate_poa_text_format(poa_type, form_data, ai_analysis)
    
    def _apply_ontario_legal_styling(self, doc):
        """Apply professional Ontario legal document styling"""
        if not DOCX_AVAILABLE:
            return
            
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(self.style_config["margins"]["top"])
            section.bottom_margin = Inches(self.style_config["margins"]["bottom"])
            section.left_margin = Inches(self.style_config["margins"]["left"])
            section.right_margin = Inches(self.style_config["margins"]["right"])
        
        # Create custom styles
        styles = doc.styles
        
        # Title style
        if 'Ontario Legal Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Ontario Legal Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = self.style_config["fonts"]["heading"]
            title_style.font.size = Pt(self.style_config["font_sizes"]["title"])
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading style
        if 'Ontario Legal Heading' not in [s.name for s in styles]:
            heading_style = styles.add_style('Ontario Legal Heading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = self.style_config["fonts"]["heading"]
            heading_style.font.size = Pt(self.style_config["font_sizes"]["heading"])
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        
        # Body style
        if 'Ontario Legal Body' not in [s.name for s in styles]:
            body_style = styles.add_style('Ontario Legal Body', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = self.style_config["fonts"]["main"]
            body_style.font.size = Pt(self.style_config["font_sizes"]["body"])
            body_style.paragraph_format.line_spacing = self.style_config["line_spacing"]
            body_style.paragraph_format.space_after = Pt(6)
    
    def _generate_will_content(self, form_data: Dict[str, Any], 
                             ai_analysis: Optional[Any]) -> Dict[str, str]:
        """Generate will content with AI enhancements"""
        content = {}
        
        # Personal information
        personal_info = form_data.get("personal_info", {})
        full_name = personal_info.get("full_name", "[TESTATOR NAME]")
        address = personal_info.get("address", "[TESTATOR ADDRESS]")
        
        # AI-enhanced opening with legal template
        opening = self.legal_kb.get_template(
            "will_opening",
            full_name=full_name,
            address=address
        )
        
        # Enhance with AI if available
        if ai_analysis:
            enhancement = self.ai_service.enhance_document_section("will", "opening", opening)
            if enhancement.confidence > 0.7:
                opening = enhancement.suggested_text
        
        content["opening"] = opening
        
        # Executor appointment
        executor_info = form_data.get("executor", {})
        executor_clause = self.legal_kb.get_template(
            "executor_clause",
            executor_name=executor_info.get("name", "[EXECUTOR NAME]"),
            executor_address=executor_info.get("address", "[EXECUTOR ADDRESS]"),
            alternate_executor=form_data.get("alternate_executor", {}).get("name", "[ALTERNATE EXECUTOR]"),
            alternate_address=form_data.get("alternate_executor", {}).get("address", "[ALTERNATE ADDRESS]")
        )
        content["executor"] = executor_clause
        
        # Guardian appointment (if applicable)
        if form_data.get("guardian"):
            guardian_info = form_data.get("guardian", {})
            guardian_clause = self.legal_kb.get_template(
                "guardian_clause",
                guardian_name=guardian_info.get("name", "[GUARDIAN NAME]"),
                guardian_address=guardian_info.get("address", "[GUARDIAN ADDRESS]"),
                alternate_guardian=form_data.get("alternate_guardian", {}).get("name", "[ALTERNATE GUARDIAN]")
            )
            content["guardian"] = guardian_clause
        
        # Beneficiaries and bequests
        content["bequests"] = self._generate_bequest_clauses(form_data.get("bequests", []))
        
        # Residuary clause
        residuary = form_data.get("residuary", {})
        content["residuary"] = f"""
RESIDUARY CLAUSE

I give, devise and bequeath all the rest, residue and remainder of my property of every nature and kind and wheresoever situate to {residuary.get("beneficiary", "[RESIDUARY BENEFICIARY]")}.
        """.strip()
        
        # Execution clause
        content["execution"] = self._generate_execution_clause()
        
        return content
    
    def _generate_poa_content(self, poa_type: str, form_data: Dict[str, Any],
                            ai_analysis: Optional[Any]) -> Dict[str, str]:
        """Generate POA content with AI enhancements"""
        content = {}
        
        # Grantor information
        grantor_info = form_data.get("grantor_info", {})
        attorney_info = form_data.get("attorney", {})
        
        # AI-enhanced opening
        template_name = f"poa_{poa_type.split('_')[1]}_opening"
        opening = self.legal_kb.get_template(
            template_name,
            grantor_name=grantor_info.get("full_name", "[GRANTOR NAME]"),
            grantor_address=grantor_info.get("address", "[GRANTOR ADDRESS]"),
            attorney_name=attorney_info.get("name", "[ATTORNEY NAME]"),
            attorney_address=attorney_info.get("address", "[ATTORNEY ADDRESS]")
        )
        content["opening"] = opening
        
        # Powers and authorities
        if poa_type == "poa_property":
            content["powers"] = self._generate_property_powers(form_data.get("powers", []))
        elif poa_type == "poa_care":
            content["care_instructions"] = self._generate_care_instructions(form_data.get("care_instructions", {}))
        
        # Conditions and restrictions
        content["conditions"] = self._generate_poa_conditions(form_data.get("conditions", []))
        
        # Execution clause
        content["execution"] = self._generate_poa_execution_clause()
        
        return content
    
    def _generate_bequest_clauses(self, bequests: List[Dict[str, Any]]) -> str:
        """Generate specific bequest clauses"""
        if not bequests:
            return ""
        
        clauses = ["SPECIFIC BEQUESTS"]
        
        for i, bequest in enumerate(bequests, 1):
            item = bequest.get("item", "[ITEM]")
            beneficiary = bequest.get("beneficiary", "[BENEFICIARY]")
            clause = f"{i}. I give and bequeath {item} to {beneficiary}."
            clauses.append(clause)
        
        return "\n\n".join(clauses)
    
    def _generate_property_powers(self, powers: List[str]) -> str:
        """Generate property power clauses"""
        default_powers = [
            "To manage and deal with all my real and personal property",
            "To buy, sell, mortgage, lease or otherwise deal with real estate",
            "To operate or dispose of any business",
            "To make gifts in accordance with a pattern of giving",
            "To employ professional assistance as necessary"
        ]
        
        power_list = powers if powers else default_powers
        
        clauses = ["POWERS AND AUTHORITIES", "I grant to my Attorney the following powers:"]
        
        for i, power in enumerate(power_list, 1):
            clauses.append(f"{i}. {power};")
        
        clauses.append("\nand generally to do all acts and things in relation to my property that I could do if personally present.")
        
        return "\n\n".join(clauses)
    
    def _generate_care_instructions(self, care_instructions: Dict[str, Any]) -> str:
        """Generate personal care instruction clauses"""
        clauses = ["PERSONAL CARE INSTRUCTIONS"]
        
        if care_instructions.get("medical_preferences"):
            clauses.append("MEDICAL TREATMENT PREFERENCES:")
            clauses.append(care_instructions["medical_preferences"])
        
        if care_instructions.get("living_arrangements"):
            clauses.append("LIVING ARRANGEMENTS:")
            clauses.append(care_instructions["living_arrangements"])
        
        if care_instructions.get("end_of_life"):
            clauses.append("END OF LIFE CARE:")
            clauses.append(care_instructions["end_of_life"])
        
        if not care_instructions:
            clauses.append("I authorize my Attorney to make decisions regarding my personal care, health care, housing, clothing, nutrition, hygiene, and safety.")
        
        return "\n\n".join(clauses)
    
    def _generate_poa_conditions(self, conditions: List[str]) -> str:
        """Generate POA conditions and restrictions"""
        if not conditions:
            return ""
        
        clauses = ["CONDITIONS AND RESTRICTIONS"]
        
        for i, condition in enumerate(conditions, 1):
            clauses.append(f"{i}. {condition}")
        
        return "\n\n".join(clauses)
    
    def _generate_execution_clause(self) -> str:
        """Generate will execution clause"""
        return """
EXECUTION

IN WITNESS WHEREOF I have to this my Last Will and Testament set my hand this _____ day of _____________, 20__.

_________________________________
[TESTATOR NAME], Testator

SIGNED, PUBLISHED AND DECLARED by [TESTATOR NAME] as and for his/her Last Will and Testament, in the presence of us, both present at the same time, who, at his/her request, in his/her presence, and in the presence of each other, have subscribed our names as witnesses.

_________________________________          _________________________________
Witness #1 Signature                      Witness #1 Name (Print)

_________________________________          _________________________________
Witness #1 Address                        Date

_________________________________          _________________________________
Witness #2 Signature                      Witness #2 Name (Print)

_________________________________          _________________________________
Witness #2 Address                        Date
        """.strip()
    
    def _generate_poa_execution_clause(self) -> str:
        """Generate POA execution clause"""
        return """
EXECUTION

IN WITNESS WHEREOF I have executed this Power of Attorney this _____ day of _____________, 20__.

_________________________________
[GRANTOR NAME], Grantor

SIGNED by [GRANTOR NAME] in the presence of us, both present at the same time, who, at his/her request, in his/her presence, and in the presence of each other, have subscribed our names as witnesses.

_________________________________          _________________________________
Witness #1 Signature                      Witness #1 Name (Print)

_________________________________          _________________________________
Witness #1 Address                        Date

_________________________________          _________________________________
Witness #2 Signature                      Witness #2 Name (Print)

_________________________________          _________________________________
Witness #2 Address                        Date
        """.strip()
    
    def _add_will_sections(self, doc, content: Dict[str, str], form_data: Dict[str, Any]):
        """Add will sections to document with professional formatting"""
        if not DOCX_AVAILABLE:
            return
        
        # Add title
        title = doc.add_paragraph("LAST WILL AND TESTAMENT", style='Ontario Legal Title')
        
        # Add opening
        doc.add_paragraph(content["opening"], style='Ontario Legal Body')
        
        # Add executor section
        doc.add_paragraph("APPOINTMENT OF EXECUTOR", style='Ontario Legal Heading')
        doc.add_paragraph(content["executor"], style='Ontario Legal Body')
        
        # Add guardian section if present
        if "guardian" in content:
            doc.add_paragraph("APPOINTMENT OF GUARDIAN", style='Ontario Legal Heading')
            doc.add_paragraph(content["guardian"], style='Ontario Legal Body')
        
        # Add bequests if present
        if content.get("bequests"):
            doc.add_paragraph(content["bequests"], style='Ontario Legal Body')
        
        # Add residuary clause
        doc.add_paragraph("RESIDUARY CLAUSE", style='Ontario Legal Heading')
        doc.add_paragraph(content["residuary"], style='Ontario Legal Body')
        
        # Add execution
        doc.add_paragraph("EXECUTION", style='Ontario Legal Heading')
        doc.add_paragraph(content["execution"], style='Ontario Legal Body')
    
    def _add_poa_sections(self, doc, poa_type: str, content: Dict[str, str], form_data: Dict[str, Any]):
        """Add POA sections to document with professional formatting"""
        if not DOCX_AVAILABLE:
            return
        
        # Add title
        title_text = "CONTINUING POWER OF ATTORNEY FOR PROPERTY" if poa_type == "poa_property" else "POWER OF ATTORNEY FOR PERSONAL CARE"
        doc.add_paragraph(title_text, style='Ontario Legal Title')
        
        # Add opening
        doc.add_paragraph(content["opening"], style='Ontario Legal Body')
        
        # Add type-specific sections
        if poa_type == "poa_property" and "powers" in content:
            doc.add_paragraph(content["powers"], style='Ontario Legal Body')
        elif poa_type == "poa_care" and "care_instructions" in content:
            doc.add_paragraph(content["care_instructions"], style='Ontario Legal Body')
        
        # Add conditions if present
        if content.get("conditions"):
            doc.add_paragraph(content["conditions"], style='Ontario Legal Body')
        
        # Add execution
        doc.add_paragraph(content["execution"], style='Ontario Legal Body')
    
    def _generate_will_text_format(self, form_data: Dict[str, Any], ai_analysis: Optional[Any]) -> Dict[str, Any]:
        """Generate will in text format as fallback"""
        content = self._generate_will_content(form_data, ai_analysis)
        
        text_content = f"""
{content['opening']}

APPOINTMENT OF EXECUTOR
{content['executor']}

{content.get('guardian', '')}

{content.get('bequests', '')}

{content['residuary']}

{content['execution']}
        """.strip()
        
        temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8')
        temp_file.write(text_content)
        temp_file.close()
        
        return {
            "document_path": temp_file.name,
            "document_type": "will",
            "format": "txt",
            "ai_analysis": asdict(ai_analysis) if ai_analysis else None,
            "generation_time": datetime.now().isoformat()
        }
    
    def _generate_poa_text_format(self, poa_type: str, form_data: Dict[str, Any], ai_analysis: Optional[Any]) -> Dict[str, Any]:
        """Generate POA in text format as fallback"""
        content = self._generate_poa_content(poa_type, form_data, ai_analysis)
        
        title = "CONTINUING POWER OF ATTORNEY FOR PROPERTY" if poa_type == "poa_property" else "POWER OF ATTORNEY FOR PERSONAL CARE"
        
        text_content = f"""
{title}

{content['opening']}

{content.get('powers', content.get('care_instructions', ''))}

{content.get('conditions', '')}

{content['execution']}
        """.strip()
        
        temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8')
        temp_file.write(text_content)
        temp_file.close()
        
        return {
            "document_path": temp_file.name,
            "document_type": poa_type,
            "format": "txt",
            "ai_analysis": asdict(ai_analysis) if ai_analysis else None,
            "generation_time": datetime.now().isoformat()
        }
    
    def generate_pdf_version(self, docx_path: str) -> str:
        """Convert DOCX to PDF (placeholder for actual PDF conversion)"""
        # In production, would use python-docx2pdf or similar
        pdf_path = docx_path.replace('.docx', '.pdf').replace('.txt', '.pdf')
        
        # For now, create a placeholder PDF file
        with open(pdf_path, 'w') as f:
            f.write("PDF conversion would happen here with proper libraries")
        
        return pdf_path
    
    def get_document_insights(self, document_type: str, form_data: Dict[str, Any]) -> Dict[str, Any]:
        """Get comprehensive document insights and recommendations"""
        return self.ai_service.generate_document_insights(document_type, form_data)