# backend/core/billing_documents.py
"""
Ontario Legal Billing Documents Generator
Generates invoices, timesheets, Bill of Costs, and cover letters
Compliant with Ontario Rules of Civil Procedure
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image as RLImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from typing import Dict, List, Any, Optional
from datetime import datetime
import io
import base64
import logging
from PIL import Image as PILImage

logger = logging.getLogger(__name__)


class OntarioBillingDocumentGenerator:
    """Generate professional billing documents for Ontario legal practice"""
    
    def __init__(self):
        self.firm_logo = None
        self.firm_info = {}
    
    def set_firm_logo(self, logo_data: str):
        """Set firm logo from base64 encoded string"""
        try:
            if logo_data:
                logo_bytes = base64.b64decode(logo_data)
                self.firm_logo = io.BytesIO(logo_bytes)
        except Exception as e:
            logger.error(f"Failed to set firm logo: {str(e)}")
    
    def set_firm_info(self, firm_info: Dict[str, Any]):
        """Set firm information for documents"""
        self.firm_info = firm_info
    
    async def generate_invoice(
        self,
        invoice_data: Dict[str, Any],
        time_entries: List[Dict[str, Any]],
        disbursements: List[Dict[str, Any]],
        template_config: Optional[Dict[str, Any]] = None,
        format: str = "pdf"
    ) -> bytes:
        """Generate professional invoice in PDF or DOCX format"""
        if format == "docx":
            return await self._generate_invoice_docx(invoice_data, time_entries, disbursements, template_config)
        else:
            return await self._generate_invoice_pdf(invoice_data, time_entries, disbursements, template_config)
    
    async def _generate_invoice_pdf(
        self,
        invoice_data: Dict[str, Any],
        time_entries: List[Dict[str, Any]],
        disbursements: List[Dict[str, Any]],
        template_config: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Generate invoice in PDF format using ReportLab"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#333333'),
            spaceAfter=10,
            spaceBefore=15
        )
        
        # Add logo if available
        if self.firm_logo and template_config and template_config.get("include_logo", True):
            try:
                logo_img = RLImage(self.firm_logo, width=2*inch, height=1*inch)
                story.append(logo_img)
                story.append(Spacer(1, 0.3*inch))
            except Exception as e:
                logger.warning(f"Failed to add logo to invoice: {str(e)}")
        
        # Firm information
        firm_name = self.firm_info.get("firm_name", "Law Firm")
        firm_address = self.firm_info.get("address", "")
        firm_phone = self.firm_info.get("phone", "")
        firm_email = self.firm_info.get("email", "")
        
        firm_info_text = f"<b>{firm_name}</b><br/>"
        if firm_address:
            firm_info_text += f"{firm_address}<br/>"
        if firm_phone:
            firm_info_text += f"Tel: {firm_phone}<br/>"
        if firm_email:
            firm_info_text += f"Email: {firm_email}<br/>"
        
        story.append(Paragraph(firm_info_text, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Invoice title
        story.append(Paragraph("INVOICE", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Invoice details
        invoice_details = [
            ["Invoice Number:", invoice_data.get("invoice_number", "N/A")],
            ["Invoice Date:", invoice_data.get("invoice_date", datetime.now().strftime("%Y-%m-%d"))],
            ["Matter:", invoice_data.get("matter_name", "N/A")],
            ["Client:", invoice_data.get("client_name", "N/A")],
        ]
        
        details_table = Table(invoice_details, colWidths=[2*inch, 4*inch])
        details_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(details_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Time entries section
        if time_entries and template_config and template_config.get("include_timesheet", True):
            story.append(Paragraph("LEGAL SERVICES", heading_style))
            
            time_data = [["Date", "Description", "Duration", "Rate", "Amount"]]
            for entry in time_entries:
                time_data.append([
                    entry.get("date", ""),
                    entry.get("description", ""),
                    f"{entry.get('duration_minutes', 0) / 60:.2f} hrs",
                    f"${entry.get('hourly_rate', 0):.2f}",
                    f"${entry.get('total_amount', 0):.2f}"
                ])
            
            time_table = Table(time_data, colWidths=[1*inch, 2.5*inch, 1*inch, 1*inch, 1*inch])
            time_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(time_table)
            story.append(Spacer(1, 0.2*inch))
        
        # Disbursements section
        if disbursements and template_config and template_config.get("include_disbursements", True):
            story.append(Paragraph("DISBURSEMENTS", heading_style))
            
            disb_data = [["Date", "Description", "Payee", "Amount", "HST", "Total"]]
            for disb in disbursements:
                disb_data.append([
                    disb.get("date", ""),
                    disb.get("description", ""),
                    disb.get("payee", ""),
                    f"${disb.get('amount', 0):.2f}",
                    f"${disb.get('hst_amount', 0):.2f}",
                    f"${disb.get('total_amount', 0):.2f}"
                ])
            
            disb_table = Table(disb_data, colWidths=[0.8*inch, 2*inch, 1.5*inch, 0.9*inch, 0.9*inch, 0.9*inch])
            disb_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (3, 0), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(disb_table)
            story.append(Spacer(1, 0.3*inch))
        
        # Summary section
        story.append(Paragraph("SUMMARY", heading_style))
        
        time_subtotal = invoice_data.get("time_subtotal", 0)
        disbursement_subtotal = invoice_data.get("disbursement_subtotal", 0)
        subtotal = invoice_data.get("subtotal", 0)
        hst = invoice_data.get("taxes", 0)
        total = invoice_data.get("total", 0)
        
        summary_data = [
            ["Legal Services Subtotal:", f"${time_subtotal:.2f}"],
            ["Disbursements Subtotal:", f"${disbursement_subtotal:.2f}"],
            ["Subtotal:", f"${subtotal:.2f}"],
            ["HST (13%):", f"${hst:.2f}"],
            ["<b>TOTAL DUE:</b>", f"<b>${total:.2f}</b>"]
        ]
        
        summary_table = Table(summary_data, colWidths=[4*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, -2), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('LINEABOVE', (0, -1), (-1, -1), 2, colors.black),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Payment terms
        payment_terms = invoice_data.get("payment_terms", "Payment due within 30 days")
        story.append(Paragraph(f"<i>Payment Terms: {payment_terms}</i>", styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def _generate_invoice_docx(
        self,
        invoice_data: Dict[str, Any],
        time_entries: List[Dict[str, Any]],
        disbursements: List[Dict[str, Any]],
        template_config: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Generate invoice in DOCX format"""
        doc = Document()
        
        # Add logo if available
        if self.firm_logo and template_config and template_config.get("include_logo", True):
            try:
                doc.add_picture(self.firm_logo, width=Inches(2))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f"Failed to add logo to invoice: {str(e)}")
        
        # Firm information
        firm_name = self.firm_info.get("firm_name", "Law Firm")
        firm_address = self.firm_info.get("address", "")
        firm_phone = self.firm_info.get("phone", "")
        firm_email = self.firm_info.get("email", "")
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(firm_name)
        run.bold = True
        run.font.size = Pt(14)
        
        if firm_address:
            doc.add_paragraph(firm_address, style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        if firm_phone:
            doc.add_paragraph(f"Tel: {firm_phone}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        if firm_email:
            doc.add_paragraph(f"Email: {firm_email}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Invoice title
        title = doc.add_heading('INVOICE', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Invoice details
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        details = [
            ("Invoice Number:", invoice_data.get("invoice_number", "N/A")),
            ("Invoice Date:", invoice_data.get("invoice_date", datetime.now().strftime("%Y-%m-%d"))),
            ("Matter:", invoice_data.get("matter_name", "N/A")),
            ("Client:", invoice_data.get("client_name", "N/A")),
        ]
        
        for i, (label, value) in enumerate(details):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value
        
        doc.add_paragraph()
        
        # Time entries
        if time_entries and template_config and template_config.get("include_timesheet", True):
            doc.add_heading('Legal Services', level=2)
            
            table = doc.add_table(rows=len(time_entries) + 1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            headers = ["Date", "Description", "Duration", "Rate", "Amount"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for i, entry in enumerate(time_entries):
                row = table.rows[i + 1]
                row.cells[0].text = entry.get("date", "")
                row.cells[1].text = entry.get("description", "")
                row.cells[2].text = f"{entry.get('duration_minutes', 0) / 60:.2f} hrs"
                row.cells[3].text = f"${entry.get('hourly_rate', 0):.2f}"
                row.cells[4].text = f"${entry.get('total_amount', 0):.2f}"
            
            doc.add_paragraph()
        
        # Disbursements
        if disbursements and template_config and template_config.get("include_disbursements", True):
            doc.add_heading('Disbursements', level=2)
            
            table = doc.add_table(rows=len(disbursements) + 1, cols=6)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            headers = ["Date", "Description", "Payee", "Amount", "HST", "Total"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for i, disb in enumerate(disbursements):
                row = table.rows[i + 1]
                row.cells[0].text = disb.get("date", "")
                row.cells[1].text = disb.get("description", "")
                row.cells[2].text = disb.get("payee", "")
                row.cells[3].text = f"${disb.get('amount', 0):.2f}"
                row.cells[4].text = f"${disb.get('hst_amount', 0):.2f}"
                row.cells[5].text = f"${disb.get('total_amount', 0):.2f}"
            
            doc.add_paragraph()
        
        # Summary
        doc.add_heading('Summary', level=2)
        
        time_subtotal = invoice_data.get("time_subtotal", 0)
        disbursement_subtotal = invoice_data.get("disbursement_subtotal", 0)
        subtotal = invoice_data.get("subtotal", 0)
        hst = invoice_data.get("taxes", 0)
        total = invoice_data.get("total", 0)
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_data = [
            ("Legal Services Subtotal:", f"${time_subtotal:.2f}"),
            ("Disbursements Subtotal:", f"${disbursement_subtotal:.2f}"),
            ("Subtotal:", f"${subtotal:.2f}"),
            ("HST (13%):", f"${hst:.2f}"),
            ("TOTAL DUE:", f"${total:.2f}")
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = summary_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            if i == len(summary_data) - 1:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].bold = True
        
        # Align amounts to the right
        for row in summary_table.rows:
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # Payment terms
        payment_terms = invoice_data.get("payment_terms", "Payment due within 30 days")
        p = doc.add_paragraph(f"Payment Terms: {payment_terms}")
        p.italic = True
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def generate_bill_of_costs(
        self,
        matter_data: Dict[str, Any],
        time_entries: List[Dict[str, Any]],
        disbursements: List[Dict[str, Any]]
    ) -> bytes:
        """Generate Bill of Costs compliant with Ontario Rules of Civil Procedure"""
        doc = Document()
        
        # Title
        title = doc.add_heading('BILL OF COSTS', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Court and matter information
        doc.add_heading('ONTARIO SUPERIOR COURT OF JUSTICE', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ("Court File No.:", matter_data.get("court_file_number", "N/A")),
            ("Plaintiff/Applicant:", matter_data.get("plaintiff", "N/A")),
            ("Defendant/Respondent:", matter_data.get("defendant", "N/A")),
            ("Matter:", matter_data.get("matter_name", "N/A")),
            ("Between:", f"{matter_data.get('plaintiff', 'N/A')} and {matter_data.get('defendant', 'N/A')}")
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value
        
        doc.add_paragraph()
        
        # Part I - Legal Fees
        doc.add_heading('PART I - LEGAL FEES', level=2)
        
        doc.add_paragraph(
            "The following are the fees charged for services necessarily or properly rendered in this proceeding:"
        )
        
        # Time entries in tabular format
        if time_entries:
            table = doc.add_table(rows=len(time_entries) + 2, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            headers = ["Date", "Description", "Time", "Amount"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            total_fees = 0
            for i, entry in enumerate(time_entries):
                row = table.rows[i + 1]
                row.cells[0].text = entry.get("date", "")
                row.cells[1].text = entry.get("description", "")
                row.cells[2].text = f"{entry.get('duration_minutes', 0) / 60:.2f} hrs"
                amount = entry.get("total_amount", 0)
                row.cells[3].text = f"${amount:.2f}"
                total_fees += amount
            
            # Total row
            total_row = table.rows[-1]
            total_row.cells[0].text = ""
            total_row.cells[1].text = ""
            total_row.cells[2].text = "TOTAL FEES:"
            total_row.cells[2].paragraphs[0].runs[0].bold = True
            total_row.cells[3].text = f"${total_fees:.2f}"
            total_row.cells[3].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Part II - Disbursements
        doc.add_heading('PART II - DISBURSEMENTS', level=2)
        
        doc.add_paragraph(
            "The following are the disbursements necessarily and reasonably incurred:"
        )
        
        # Disbursements in tabular format
        if disbursements:
            table = doc.add_table(rows=len(disbursements) + 2, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            headers = ["Date", "Description", "Payee", "Amount"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            total_disbursements = 0
            for i, disb in enumerate(disbursements):
                row = table.rows[i + 1]
                row.cells[0].text = disb.get("date", "")
                row.cells[1].text = disb.get("description", "")
                row.cells[2].text = disb.get("payee", "")
                amount = disb.get("total_amount", 0)  # Including HST
                row.cells[3].text = f"${amount:.2f}"
                total_disbursements += amount
            
            # Total row
            total_row = table.rows[-1]
            total_row.cells[0].text = ""
            total_row.cells[1].text = ""
            total_row.cells[2].text = "TOTAL DISBURSEMENTS:"
            total_row.cells[2].paragraphs[0].runs[0].bold = True
            total_row.cells[3].text = f"${total_disbursements:.2f}"
            total_row.cells[3].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Summary
        doc.add_heading('SUMMARY', level=2)
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        fees_subtotal = sum(entry.get("total_amount", 0) for entry in time_entries) if time_entries else 0
        disb_subtotal = sum(disb.get("total_amount", 0) for disb in disbursements) if disbursements else 0
        subtotal = fees_subtotal + disb_subtotal
        hst = fees_subtotal * 0.13  # HST only on fees
        total = subtotal + hst
        
        summary_data = [
            ("Fees (Part I):", f"${fees_subtotal:.2f}"),
            ("Disbursements (Part II):", f"${disb_subtotal:.2f}"),
            ("HST on Fees:", f"${hst:.2f}"),
            ("TOTAL:", f"${total:.2f}")
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = summary_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            if i == len(summary_data) - 1:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Certificate
        doc.add_paragraph()
        doc.add_paragraph(
            f"I certify that the attached bill of costs was assessed and allowed in the amount of ${total:.2f}."
        )
        doc.add_paragraph()
        doc.add_paragraph("Date: ___________________")
        doc.add_paragraph()
        doc.add_paragraph("_____________________________________")
        doc.add_paragraph("Assessment Officer")
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


class AIEnhancedCoverLetterGenerator:
    """AI-powered cover letter generator for invoices and documents"""
    
    def __init__(self):
        self.firm_logo = None
        self.firm_info = {}
    
    def set_firm_logo(self, logo_data: str):
        """Set firm logo from base64 encoded string"""
        try:
            if logo_data:
                logo_bytes = base64.b64decode(logo_data)
                self.firm_logo = io.BytesIO(logo_bytes)
        except Exception as e:
            logger.error(f"Failed to set firm logo: {str(e)}")
    
    def set_firm_info(self, firm_info: Dict[str, Any]):
        """Set firm information"""
        self.firm_info = firm_info
    
    async def generate_cover_letter(
        self,
        client_info: Dict[str, Any],
        documents_included: List[str],
        invoice_data: Optional[Dict[str, Any]] = None,
        additional_notes: Optional[str] = None,
        format: str = "pdf"
    ) -> bytes:
        """Generate AI-enhanced professional cover letter"""
        
        # Generate letter content using AI-style template
        letter_content = self._generate_letter_content(
            client_info, documents_included, invoice_data, additional_notes
        )
        
        if format == "docx":
            return await self._generate_letter_docx(letter_content, client_info)
        else:
            return await self._generate_letter_pdf(letter_content, client_info)
    
    def _generate_letter_content(
        self,
        client_info: Dict[str, Any],
        documents_included: List[str],
        invoice_data: Optional[Dict[str, Any]],
        additional_notes: Optional[str]
    ) -> str:
        """Generate professional letter content"""
        
        client_name = client_info.get("name", "Dear Client")
        matter_name = client_info.get("matter_name", "your matter")
        
        # Start with greeting
        content = f"Dear {client_name},\n\n"
        
        # Opening paragraph
        content += f"Re: {matter_name}\n\n"
        
        # Invoice section if applicable
        if invoice_data:
            invoice_number = invoice_data.get("invoice_number", "N/A")
            total = invoice_data.get("total", 0)
            content += f"Please find enclosed our invoice number {invoice_number} in the amount of ${total:.2f} "
            content += "for professional services rendered in connection with the above-noted matter.\n\n"
        
        # Documents section
        if documents_included:
            content += "The following documents are enclosed for your review and records:\n\n"
            for i, doc in enumerate(documents_included, 1):
                content += f"{i}. {doc}\n"
            content += "\n"
        
        # Payment instructions if invoice
        if invoice_data:
            payment_terms = invoice_data.get("payment_terms", "Payment is due within 30 days")
            content += f"{payment_terms}. "
            content += "Payment can be made by cheque, e-transfer, or direct deposit. "
            content += "Please reference the invoice number with your payment.\n\n"
        
        # Additional notes
        if additional_notes:
            content += f"{additional_notes}\n\n"
        
        # Closing
        content += "Should you have any questions or require clarification regarding any aspect of this invoice or the enclosed documents, "
        content += "please do not hesitate to contact our office.\n\n"
        content += "We appreciate your continued confidence in our services.\n\n"
        content += "Yours truly,\n\n"
        
        firm_name = self.firm_info.get("firm_name", "Law Firm")
        lawyer_name = self.firm_info.get("lawyer_name", "")
        
        content += f"{firm_name}\n"
        if lawyer_name:
            content += f"{lawyer_name}\n"
        
        return content
    
    async def _generate_letter_pdf(self, content: str, client_info: Dict[str, Any]) -> bytes:
        """Generate cover letter in PDF format"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch, bottomMargin=1*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Add logo if available
        if self.firm_logo:
            try:
                logo_img = RLImage(self.firm_logo, width=2*inch, height=1*inch)
                story.append(logo_img)
                story.append(Spacer(1, 0.3*inch))
            except Exception as e:
                logger.warning(f"Failed to add logo to letter: {str(e)}")
        
        # Firm letterhead
        firm_name = self.firm_info.get("firm_name", "Law Firm")
        firm_address = self.firm_info.get("address", "")
        firm_phone = self.firm_info.get("phone", "")
        firm_email = self.firm_info.get("email", "")
        
        letterhead_style = ParagraphStyle(
            'Letterhead',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_CENTER
        )
        
        story.append(Paragraph(f"<b>{firm_name}</b>", letterhead_style))
        if firm_address:
            story.append(Paragraph(firm_address, letterhead_style))
        if firm_phone:
            story.append(Paragraph(f"Tel: {firm_phone}", letterhead_style))
        if firm_email:
            story.append(Paragraph(f"Email: {firm_email}", letterhead_style))
        
        story.append(Spacer(1, 0.5*inch))
        
        # Date
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_LEFT
        )
        story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), date_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Client address (if available)
        if client_info.get("address"):
            story.append(Paragraph(client_info.get("name", ""), styles['Normal']))
            story.append(Paragraph(client_info["address"], styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
        
        # Letter content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph.replace('\n', '<br/>'), styles['Normal']))
                story.append(Spacer(1, 0.15*inch))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def _generate_letter_docx(self, content: str, client_info: Dict[str, Any]) -> bytes:
        """Generate cover letter in DOCX format"""
        doc = Document()
        
        # Add logo if available
        if self.firm_logo:
            try:
                doc.add_picture(self.firm_logo, width=Inches(2))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f"Failed to add logo to letter: {str(e)}")
        
        # Firm letterhead
        firm_name = self.firm_info.get("firm_name", "Law Firm")
        firm_address = self.firm_info.get("address", "")
        firm_phone = self.firm_info.get("phone", "")
        firm_email = self.firm_info.get("email", "")
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(firm_name)
        run.bold = True
        run.font.size = Pt(12)
        
        if firm_address:
            doc.add_paragraph(firm_address, style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        if firm_phone:
            doc.add_paragraph(f"Tel: {firm_phone}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        if firm_email:
            doc.add_paragraph(f"Email: {firm_email}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Date
        doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        doc.add_paragraph()
        
        # Client address (if available)
        if client_info.get("address"):
            doc.add_paragraph(client_info.get("name", ""))
            doc.add_paragraph(client_info["address"])
            doc.add_paragraph()
        
        # Letter content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
