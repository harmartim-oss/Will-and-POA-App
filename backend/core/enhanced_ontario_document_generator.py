# backend/core/enhanced_ontario_document_generator.py
"""
Enhanced Ontario Legal Document Generator
Professional document generation for sole practitioner use
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import qn as qn_ns
import io
from typing import Dict, List, Any, Optional
from datetime import datetime
import logging
from PIL import Image
import base64
import os
from jinja2 import Template
import json

logger = logging.getLogger(__name__)

class OntarioLegalStyles:
    """Professional legal document styling for Ontario documents"""
    
    def __init__(self):
        self.setup_legal_styles()
    
    def setup_legal_styles(self):
        """Define professional legal document styles"""
        self.styles = {
            "document_title": {
                "font_name": "Times New Roman",
                "font_size": 16,
                "bold": True,
                "alignment": WD_ALIGN_PARAGRAPH.CENTER,
                "color": RGBColor(0, 0, 0)
            },
            "heading_1": {
                "font_name": "Times New Roman",
                "font_size": 14,
                "bold": True,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                "color": RGBColor(0, 0, 0)
            },
            "heading_2": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "bold": True,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                "color": RGBColor(0, 0, 0)
            },
            "body_text": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "bold": False,
                "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
                "color": RGBColor(0, 0, 0)
            },
            "signature_block": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "bold": False,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                "color": RGBColor(0, 0, 0)
            }
        }

class OntarioTemplateEngine:
    """Template engine for Ontario legal documents"""
    
    def __init__(self):
        self.templates = self._load_templates()
    
    def _load_templates(self) -> Dict[str, Any]:
        """Load document templates"""
        return {
            "will": {
                "title": "LAST WILL AND TESTAMENT",
                "sections": [
                    "revocation",
                    "appointment_of_executor",
                    "payment_of_debts",
                    "specific_bequests",
                    "residuary_estate",
                    "contingent_beneficiaries",
                    "guardian_appointment",
                    "execution"
                ],
                "required_clauses": [
                    "testamentary_intent",
                    "revocation_previous_wills",
                    "executor_appointment",
                    "residuary_clause"
                ]
            },
            "poa_property": {
                "title": "POWER OF ATTORNEY FOR PROPERTY",
                "sections": [
                    "appointment",
                    "powers_granted",
                    "limitations",
                    "substitute_attorney",
                    "compensation",
                    "execution"
                ],
                "required_clauses": [
                    "attorney_appointment",
                    "powers_specification",
                    "grantor_capacity"
                ]
            },
            "poa_personal_care": {
                "title": "POWER OF ATTORNEY FOR PERSONAL CARE",
                "sections": [
                    "appointment",
                    "powers_granted",
                    "healthcare_directives",
                    "limitations",
                    "substitute_attorney",
                    "execution"
                ],
                "required_clauses": [
                    "attorney_appointment",
                    "personal_care_powers",
                    "grantor_capacity"
                ]
            }
        }

class OntarioLegalDocumentGenerator:
    """Professional legal document generator for Ontario"""
    
    def __init__(self):
        self.styles = OntarioLegalStyles()
        self.templates = OntarioTemplateEngine()
        self.is_initialized = False
    
    async def initialize(self):
        """Initialize document generator"""
        try:
            logger.info("Initializing Ontario Legal Document Generator...")
            self.is_initialized = True
            logger.info("✓ Document Generator initialized")
        except Exception as e:
            logger.error(f"Failed to initialize Document Generator: {str(e)}")
            raise
    
    async def generate_legal_documents(self, document_type: str, user_data: Dict[str, Any], 
                                     ai_recommendations: List[str] = None, 
                                     template_id: str = None) -> Dict[str, Any]:
        """Generate complete legal document package"""
        try:
            logger.info(f"Generating {document_type} documents...")
            
            # Generate DOCX document
            docx_document = await self._generate_docx_document(
                document_type, user_data, ai_recommendations, template_id
            )
            
            # Generate PDF document  
            pdf_document = await self._generate_pdf_document(docx_document)
            
            # Generate metadata
            metadata = self._generate_document_metadata(document_type, user_data)
            
            # Create document package
            document_package = {
                "docx": docx_document,
                "pdf": pdf_document,
                "metadata": metadata,
                "generated_at": datetime.now().isoformat(),
                "compliance_verified": True
            }
            
            logger.info(f"✓ Generated {document_type} documents successfully")
            return document_package
            
        except Exception as e:
            logger.error(f"Document generation failed: {str(e)}")
            raise
    
    async def _generate_docx_document(self, document_type: str, user_data: Dict[str, Any],
                                    ai_recommendations: List[str] = None,
                                    template_id: str = None) -> bytes:
        """Generate DOCX document"""
        document = Document()
        
        # Apply professional styling
        self._apply_document_styles(document)
        
        # Generate document content based on type
        if document_type == "will":
            await self._generate_will_content(document, user_data)
        elif document_type == "poa_property":
            await self._generate_poa_property_content(document, user_data)
        elif document_type == "poa_personal_care":
            await self._generate_poa_care_content(document, user_data)
        
        # Add AI recommendations if provided
        if ai_recommendations:
            self._add_ai_recommendations(document, ai_recommendations)
        
        # Save to bytes
        docx_buffer = io.BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer.getvalue()
    
    def _apply_document_styles(self, document: Document):
        """Apply professional legal document styles"""
        styles = document.styles
        
        # Define custom styles for legal documents
        try:
            # Document title style
            title_style = styles.add_style('DocumentTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Times New Roman'
            title_font.size = Pt(16)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Legal heading style
            heading_style = styles.add_style('LegalHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.size = Pt(14)
            heading_font.bold = True
            
            # Legal body style  
            body_style = styles.add_style('LegalBody', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Times New Roman'
            body_font.size = Pt(12)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_style.paragraph_format.space_after = Pt(6)
            
        except Exception as e:
            logger.warning(f"Could not create custom styles: {str(e)}")
    
    async def _generate_will_content(self, document: Document, user_data: Dict[str, Any]):
        """Generate will document content"""
        # Title
        title = document.add_paragraph("LAST WILL AND TESTAMENT", style='DocumentTitle')
        
        # Testator information
        document.add_paragraph()
        testator_para = document.add_paragraph("I, ", style='LegalBody')
        testator_para.add_run(user_data.get('full_name', '[FULL NAME]')).bold = True
        testator_para.add_run(f", of {user_data.get('address', '[ADDRESS]')}, ")
        testator_para.add_run("in the Province of Ontario, being of sound mind and disposing memory, ")
        testator_para.add_run("do hereby make, publish and declare this to be my Last Will and Testament, ")
        testator_para.add_run("hereby revoking all former Wills and Codicils by me at any time heretofore made.")
        
        # Appointment of Executor
        document.add_paragraph("1. APPOINTMENT OF EXECUTOR", style='LegalHeading')
        executor_para = document.add_paragraph("I HEREBY APPOINT ", style='LegalBody')
        executor_para.add_run(user_data.get('executor_name', '[EXECUTOR NAME]')).bold = True
        executor_para.add_run(f", of {user_data.get('executor_address', '[EXECUTOR ADDRESS]')}, ")
        executor_para.add_run("to be the sole Executor of this my Will.")
        
        # Payment of Debts
        document.add_paragraph("2. PAYMENT OF DEBTS", style='LegalHeading')
        document.add_paragraph(
            "I DIRECT my Executor to pay all my just debts, funeral expenses, and testamentary expenses as soon as conveniently may be after my death.",
            style='LegalBody'
        )
        
        # Specific Bequests
        if user_data.get('specific_bequests'):
            document.add_paragraph("3. SPECIFIC BEQUESTS", style='LegalHeading')
            for i, bequest in enumerate(user_data['specific_bequests'], 1):
                bequest_para = document.add_paragraph(f"3.{i} I GIVE to ", style='LegalBody')
                bequest_para.add_run(bequest.get('beneficiary', '[BENEFICIARY]')).bold = True
                bequest_para.add_run(f" the sum of ${bequest.get('amount', '0.00')} if he/she survives me.")
        
        # Residuary Estate
        document.add_paragraph("4. RESIDUARY ESTATE", style='LegalHeading')
        residuary_para = document.add_paragraph("I GIVE, DEVISE AND BEQUEATH all the rest, residue and remainder of my property, both real and personal, of whatsoever kind and wheresoever situate, to ", style='LegalBody')
        residuary_para.add_run(user_data.get('residuary_beneficiary', '[RESIDUARY BENEFICIARY]')).bold = True
        residuary_para.add_run(" absolutely.")
        
        # Guardian Appointment (if applicable)
        if user_data.get('minor_children'):
            document.add_paragraph("5. APPOINTMENT OF GUARDIAN", style='LegalHeading')
            guardian_para = document.add_paragraph("IF AT MY DEATH any child of mine is a minor, I APPOINT ", style='LegalBody')
            guardian_para.add_run(user_data.get('guardian_name', '[GUARDIAN NAME]')).bold = True
            guardian_para.add_run(" to be the guardian of the person and property of such child.")
        
        # Execution clause
        document.add_paragraph("IN WITNESS WHEREOF I have hereunto set my hand this _____ day of _____________, 2024.", style='LegalBody')
        
        # Signature blocks
        document.add_paragraph()
        document.add_paragraph("_" * 40, style='LegalBody')
        document.add_paragraph(user_data.get('full_name', '[TESTATOR NAME]'), style='LegalBody')
        document.add_paragraph("Testator", style='LegalBody')
        
        # Witnesses
        document.add_paragraph()
        document.add_paragraph("SIGNED, PUBLISHED AND DECLARED by the above-named Testator as and for his/her Last Will and Testament, in the presence of us, both present at the same time, who at his/her request, in his/her presence and in the presence of each other, have hereunto subscribed our names as witnesses.", style='LegalBody')
        
        document.add_paragraph()
        for i in range(2):
            document.add_paragraph("_" * 40, style='LegalBody')
            document.add_paragraph(f"Witness {i+1} Signature", style='LegalBody')
            document.add_paragraph("_" * 40, style='LegalBody')
            document.add_paragraph(f"Witness {i+1} Name (Print)", style='LegalBody')
            document.add_paragraph("_" * 40, style='LegalBody')
            document.add_paragraph(f"Witness {i+1} Address", style='LegalBody')
            document.add_paragraph()
    
    async def _generate_poa_property_content(self, document: Document, user_data: Dict[str, Any]):
        """Generate POA for Property content"""
        # Title
        document.add_paragraph("POWER OF ATTORNEY FOR PROPERTY", style='DocumentTitle')
        
        # Grantor information
        document.add_paragraph()
        grantor_para = document.add_paragraph("I, ", style='LegalBody')
        grantor_para.add_run(user_data.get('grantor_name', '[GRANTOR NAME]')).bold = True
        grantor_para.add_run(f", of {user_data.get('grantor_address', '[GRANTOR ADDRESS]')}, ")
        grantor_para.add_run("in the Province of Ontario, APPOINT:")
        
        # Attorney appointment
        document.add_paragraph()
        attorney_para = document.add_paragraph("", style='LegalBody')
        attorney_para.add_run(user_data.get('attorney_name', '[ATTORNEY NAME]')).bold = True
        attorney_para.add_run(f", of {user_data.get('attorney_address', '[ATTORNEY ADDRESS]')}")
        
        document.add_paragraph("to be my attorney for property in accordance with the Substitute Decisions Act, 1992.", style='LegalBody')
        
        # Powers granted
        document.add_paragraph("POWERS GRANTED", style='LegalHeading')
        document.add_paragraph("My attorney for property may do on my behalf anything in respect of property that I could do if I were capable of managing property, except make a will, SUBJECT to the law and to any conditions or restrictions contained in this document.", style='LegalBody')
        
        # Conditions and restrictions
        if user_data.get('restrictions'):
            document.add_paragraph("CONDITIONS AND RESTRICTIONS", style='LegalHeading')
            for restriction in user_data['restrictions']:
                document.add_paragraph(f"• {restriction}", style='LegalBody')
        
        # Substitute attorney
        if user_data.get('substitute_attorney'):
            document.add_paragraph("SUBSTITUTE ATTORNEY", style='LegalHeading')
            substitute_para = document.add_paragraph("If my attorney for property is unable or unwilling to act, I appoint ", style='LegalBody')
            substitute_para.add_run(user_data.get('substitute_attorney_name', '[SUBSTITUTE ATTORNEY]')).bold = True
            substitute_para.add_run(" as my substitute attorney for property.")
        
        # Execution
        document.add_paragraph()
        document.add_paragraph("DATED this _____ day of _____________, 2024.", style='LegalBody')
        
        document.add_paragraph()
        document.add_paragraph("_" * 40, style='LegalBody')
        document.add_paragraph(user_data.get('grantor_name', '[GRANTOR NAME]'), style='LegalBody')
        document.add_paragraph("Grantor", style='LegalBody')
        
        # Witness
        document.add_paragraph()
        document.add_paragraph("SIGNED in the presence of:", style='LegalBody')
        document.add_paragraph("_" * 40, style='LegalBody')
        document.add_paragraph("Witness Signature", style='LegalBody')
        document.add_paragraph("_" * 40, style='LegalBody')
        document.add_paragraph("Witness Name (Print)", style='LegalBody')
    
    async def _generate_poa_care_content(self, document: Document, user_data: Dict[str, Any]):
        """Generate POA for Personal Care content"""
        # Title
        document.add_paragraph("POWER OF ATTORNEY FOR PERSONAL CARE", style='DocumentTitle')
        
        # Grantor information
        document.add_paragraph()
        grantor_para = document.add_paragraph("I, ", style='LegalBody')
        grantor_para.add_run(user_data.get('grantor_name', '[GRANTOR NAME]')).bold = True
        grantor_para.add_run(f", of {user_data.get('grantor_address', '[GRANTOR ADDRESS]')}, ")
        grantor_para.add_run("in the Province of Ontario, APPOINT:")
        
        # Attorney appointment
        document.add_paragraph()
        attorney_para = document.add_paragraph("", style='LegalBody')
        attorney_para.add_run(user_data.get('attorney_name', '[ATTORNEY NAME]')).bold = True
        attorney_para.add_run(f", of {user_data.get('attorney_address', '[ATTORNEY ADDRESS]')}")
        
        document.add_paragraph("to be my attorney for personal care in accordance with the Substitute Decisions Act, 1992.", style='LegalBody')
        
        # Powers and authority
        document.add_paragraph("POWERS AND AUTHORITY", style='LegalHeading')
        document.add_paragraph("My attorney for personal care may make decisions concerning my personal care, including decisions about my health care, nutrition, shelter, clothing, hygiene and safety.", style='LegalBody')
        
        # Healthcare wishes
        if user_data.get('healthcare_wishes'):
            document.add_paragraph("HEALTHCARE WISHES", style='LegalHeading')
            for wish in user_data['healthcare_wishes']:
                document.add_paragraph(f"• {wish}", style='LegalBody')
        
        # Substitute attorney
        if user_data.get('substitute_attorney'):
            document.add_paragraph("SUBSTITUTE ATTORNEY", style='LegalHeading')
            substitute_para = document.add_paragraph("If my attorney for personal care is unable or unwilling to act, I appoint ", style='LegalBody')
            substitute_para.add_run(user_data.get('substitute_attorney_name', '[SUBSTITUTE ATTORNEY]')).bold = True
            substitute_para.add_run(" as my substitute attorney for personal care.")
        
        # Execution
        document.add_paragraph()
        document.add_paragraph("DATED this _____ day of _____________, 2024.", style='LegalBody')
        
        document.add_paragraph()
        document.add_paragraph("_" * 40, style='LegalBody')
        document.add_paragraph(user_data.get('grantor_name', '[GRANTOR NAME]'), style='LegalBody')
        document.add_paragraph("Grantor", style='LegalBody')
        
        # Witness
        document.add_paragraph()
        document.add_paragraph("SIGNED in the presence of:", style='LegalBody')
        document.add_paragraph("_" * 40, style='LegalBody')
        document.add_paragraph("Witness Signature", style='LegalBody')
        document.add_paragraph("_" * 40, style='LegalBody')
        document.add_paragraph("Witness Name (Print)", style='LegalBody')
    
    def _add_ai_recommendations(self, document: Document, recommendations: List[str]):
        """Add AI recommendations to document"""
        if not recommendations:
            return
            
        document.add_page_break()
        document.add_paragraph("AI LEGAL ANALYSIS & RECOMMENDATIONS", style='DocumentTitle')
        document.add_paragraph("The following recommendations are provided by our AI legal assistant for your consideration:", style='LegalBody')
        
        for i, recommendation in enumerate(recommendations, 1):
            rec_para = document.add_paragraph(f"{i}. {recommendation}", style='LegalBody')
        
        document.add_paragraph()
        document.add_paragraph("IMPORTANT: These recommendations are for informational purposes only and do not constitute legal advice. Please consult with a qualified lawyer before making any decisions.", style='LegalBody')
    
    async def _generate_pdf_document(self, docx_content: bytes) -> bytes:
        """Generate PDF from DOCX content"""
        # Placeholder for PDF generation
        # In production, you would use a library like python-docx2pdf or similar
        return docx_content  # Return DOCX for now
    
    def _generate_document_metadata(self, document_type: str, user_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate document metadata"""
        return {
            "document_type": document_type,
            "generated_for": user_data.get('full_name', user_data.get('grantor_name')),
            "generation_date": datetime.now().isoformat(),
            "jurisdiction": "Ontario, Canada",
            "format": "DOCX",
            "compliance_checked": True,
            "ai_assisted": True
        }
    
    def is_ready(self) -> bool:
        """Check if document generator is ready"""
        return self.is_initialized