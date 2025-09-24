"""
AI-Enhanced Document Generator for Ontario Legal Documents
Implements professional document generation with AI-powered content optimization
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

import os
import json
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
from jinja2 import Template, Environment, FileSystemLoader
import asyncio

from .enhanced_ai_legal_service import EnhancedAILegalService
from .legal_knowledge import OntarioLegalKnowledgeBase

logger = logging.getLogger(__name__)

class AIEnhancedDocumentGenerator:
    """AI-Enhanced document generator with professional Ontario legal styling"""
    
    def __init__(self):
        self.ai_service = EnhancedAILegalService()
        self.legal_kb = OntarioLegalKnowledgeBase()
        self.template_engine = None
        self.is_initialized = False
        
        # Ontario legal styling configuration
        self.ontario_styles = {
            "document_title": {
                "font_name": "Times New Roman",
                "font_size": 16,
                "bold": True,
                "alignment": "center",
                "spacing_after": 24
            },
            "section_heading": {
                "font_name": "Times New Roman", 
                "font_size": 14,
                "bold": True,
                "spacing_before": 18,
                "spacing_after": 12
            },
            "body_text": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "line_spacing": 1.15,
                "spacing_after": 6
            },
            "legal_clause": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "indent": 0.5,
                "spacing_after": 12
            },
            "signature_block": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "spacing_before": 36,
                "spacing_after": 12
            }
        }
        
        # Document templates directory
        self.templates_dir = os.path.join(os.path.dirname(__file__), "templates")
        
    async def initialize(self):
        """Initialize the document generator"""
        try:
            logger.info("Initializing AI-Enhanced Document Generator...")
            
            # Initialize AI service
            await self.ai_service.initialize()
            
            # Initialize legal knowledge base
            await self.legal_kb.initialize()
            
            # Setup Jinja2 template environment
            if os.path.exists(self.templates_dir):
                self.template_engine = Environment(
                    loader=FileSystemLoader(self.templates_dir),
                    trim_blocks=True,
                    lstrip_blocks=True
                )
            else:
                # Create templates directory and basic templates
                os.makedirs(self.templates_dir, exist_ok=True)
                self._create_default_templates()
                self.template_engine = Environment(
                    loader=FileSystemLoader(self.templates_dir),
                    trim_blocks=True,
                    lstrip_blocks=True
                )
            
            self.is_initialized = True
            logger.info("AI-Enhanced Document Generator initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize document generator: {str(e)}")
            self.is_initialized = False
    
    def _create_default_templates(self):
        """Create default Ontario legal document templates"""
        
        # Ontario Will Template
        will_template = """
LAST WILL AND TESTAMENT

I, {{ testator.full_name }}, of {{ testator.address }}, {{ testator.city }}, Ontario, {{ testator.postal_code }}, 
being of sound mind and disposing memory, do hereby make, publish and declare this to be my Last Will and Testament, 
hereby revoking all former Wills and Codicils by me at any time heretofore made.

1. APPOINTMENT OF EXECUTOR{% if executors|length > 1 %}S{% endif %}
{% for executor in executors %}
I APPOINT {{ executor.name }} of {{ executor.address }} to be {% if executors|length > 1 %}an {% endif %}Executor{% if executors|length > 1 %} and Trustee{% endif %} of this my Will.
{% endfor %}

{% if alternate_executors %}
If any of the above-named Executor{% if executors|length > 1 %}s{% endif %} shall predecease me or be unable or unwilling to act, 
I APPOINT the following as alternate Executor{% if alternate_executors|length > 1 %}s{% endif %}:
{% for alt_executor in alternate_executors %}
{{ alt_executor.name }} of {{ alt_executor.address }}
{% endfor %}
{% endif %}

2. PAYMENT OF DEBTS
I DIRECT my Executor{% if executors|length > 1 %}s{% endif %} to pay my just debts, funeral expenses, and testamentary expenses.

3. SPECIFIC BEQUESTS
{% if specific_bequests %}
{% for bequest in specific_bequests %}
I GIVE, DEVISE AND BEQUEATH {{ bequest.description }} to {{ bequest.beneficiary }}.
{% endfor %}
{% endif %}

4. RESIDUARY CLAUSE
I GIVE, DEVISE AND BEQUEATH all the rest, residue and remainder of my estate, both real and personal, 
wheresoever situate, unto {{ residuary_beneficiary }} absolutely.

{% if guardian_appointment and minors %}
5. GUARDIAN APPOINTMENT
I APPOINT {{ guardian_appointment.name }} of {{ guardian_appointment.address }} to be the Guardian of the person{% if guardian_appointment.property %} and property{% endif %} 
of any child of mine who may be under the age of eighteen (18) years at the time of my death.
{% endif %}

IN WITNESS WHEREOF I have to this my Last Will and Testament, written on this and the preceding page, 
subscribed my name this _____ day of ____________, {{ current_year }}.

SIGNED, PUBLISHED AND DECLARED by the said {{ testator.full_name }} as and for {{ testator.pronoun_possessive }} Last Will and Testament, 
in the presence of us, both present at the same time, who, at {{ testator.pronoun_possessive }} request, in {{ testator.pronoun_possessive }} presence, 
and in the presence of each other, have hereunto subscribed our names as witnesses.

_________________________________         _________________________________
{{ testator.full_name }}                    Date
Testator

WITNESSES:

_________________________________         _________________________________
Witness Signature                         Date
{{ witness1.name }}
{{ witness1.address }}

_________________________________         _________________________________  
Witness Signature                         Date
{{ witness2.name }}
{{ witness2.address }}
"""
        
        # Ontario Power of Attorney for Property Template
        poa_property_template = """
POWER OF ATTORNEY FOR PROPERTY

I, {{ grantor.full_name }}, of {{ grantor.address }}, {{ grantor.city }}, Ontario, {{ grantor.postal_code }}, 
APPOINT {{ attorney.name }} of {{ attorney.address }} to be my attorney for property.

{% if alternate_attorney %}
If my attorney is unable or unwilling to act as my attorney, I APPOINT {{ alternate_attorney.name }} 
of {{ alternate_attorney.address }} to be my attorney for property in place of my first-named attorney.
{% endif %}

I GIVE my attorney the authority to do on my behalf anything in respect of property that I could do if capable of managing property, 
except make a will, subject to the law and to any conditions or restrictions contained in this document.

CONDITIONS AND RESTRICTIONS:
{% if conditions %}
{% for condition in conditions %}
{{ loop.index }}. {{ condition }}
{% endfor %}
{% else %}
None.
{% endif %}

{% if powers_commence %}
COMMENCEMENT:
This Power of Attorney for Property shall commence {{ powers_commence }}.
{% endif %}

{% if compensation %}
COMPENSATION:
My attorney shall be entitled to compensation for services as my attorney as provided by law.
{% endif %}

SIGNED this _____ day of ____________, {{ current_year }}.

_________________________________
{{ grantor.full_name }}
Grantor

WITNESS:

_________________________________         _________________________________
Witness Signature                         Date
{{ witness.name }}
{{ witness.address }}
"""
        
        # Save templates
        with open(os.path.join(self.templates_dir, "ontario_will.txt"), "w") as f:
            f.write(will_template.strip())
            
        with open(os.path.join(self.templates_dir, "ontario_poa_property.txt"), "w") as f:
            f.write(poa_property_template.strip())

    async def generate_ontario_will(self, user_data: Dict[str, Any], ai_enhancements: bool = True) -> Dict[str, Any]:
        """Generate Ontario-compliant Will with AI enhancements"""
        try:
            logger.info("Generating Ontario Will document...")
            
            # AI-enhanced content optimization
            if ai_enhancements and self.ai_service.is_ready():
                user_data = await self._ai_enhance_content(user_data, "will")
            
            # Get AI analysis and recommendations
            analysis = None
            if self.ai_service.is_ready():
                analysis = self.ai_service.analyze_document("will", user_data)
            
            # Generate document content
            template = self.template_engine.get_template("ontario_will.txt")
            
            # Prepare template context
            context = self._prepare_will_context(user_data)
            content = template.render(**context)
            
            # Generate DOCX document
            docx_path = None
            if DOCX_AVAILABLE:
                docx_path = await self._generate_docx_will(content, context)
            
            # Generate PDF document  
            pdf_path = None
            if REPORTLAB_AVAILABLE or WEASYPRINT_AVAILABLE:
                pdf_path = await self._generate_pdf_document(content, "will")
            
            return {
                "success": True,
                "document_type": "will",
                "content": content,
                "docx_path": docx_path,
                "pdf_path": pdf_path,
                "ai_analysis": asdict(analysis) if analysis else None,
                "metadata": {
                    "generated_at": datetime.now().isoformat(),
                    "ai_enhanced": ai_enhancements,
                    "compliance_score": analysis.compliance_score if analysis else None
                }
            }
            
        except Exception as e:
            logger.error(f"Will generation failed: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "document_type": "will"
            }

    async def generate_ontario_poa(self, user_data: Dict[str, Any], poa_type: str, ai_enhancements: bool = True) -> Dict[str, Any]:
        """Generate Ontario-compliant Power of Attorney with AI enhancements"""
        try:
            logger.info(f"Generating Ontario POA ({poa_type}) document...")
            
            # AI-enhanced content optimization
            if ai_enhancements and self.ai_service.is_ready():
                user_data = await self._ai_enhance_content(user_data, poa_type)
            
            # Get AI analysis
            analysis = None
            if self.ai_service.is_ready():
                analysis = self.ai_service.analyze_document(poa_type, user_data)
            
            # Generate document content
            template_name = f"ontario_poa_{poa_type.split('_')[-1]}.txt"
            template = self.template_engine.get_template(template_name)
            
            # Prepare template context
            context = self._prepare_poa_context(user_data, poa_type)
            content = template.render(**context)
            
            # Generate DOCX document
            docx_path = None
            if DOCX_AVAILABLE:
                docx_path = await self._generate_docx_poa(content, context, poa_type)
            
            # Generate PDF document
            pdf_path = None  
            if REPORTLAB_AVAILABLE or WEASYPRINT_AVAILABLE:
                pdf_path = await self._generate_pdf_document(content, poa_type)
            
            return {
                "success": True,
                "document_type": poa_type,
                "content": content,
                "docx_path": docx_path,
                "pdf_path": pdf_path,
                "ai_analysis": asdict(analysis) if analysis else None,
                "metadata": {
                    "generated_at": datetime.now().isoformat(),
                    "ai_enhanced": ai_enhancements,
                    "compliance_score": analysis.compliance_score if analysis else None
                }
            }
            
        except Exception as e:
            logger.error(f"POA generation failed: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "document_type": poa_type
            }

    async def _ai_enhance_content(self, user_data: Dict[str, Any], document_type: str) -> Dict[str, Any]:
        """Use AI to enhance and optimize document content"""
        try:
            # Get AI recommendations
            insights = self.ai_service.generate_document_insights(document_type, user_data)
            
            # Apply AI recommendations to content
            enhanced_data = user_data.copy()
            
            # Add AI-suggested improvements
            if insights.get("key_recommendations"):
                enhanced_data["ai_recommendations"] = insights["key_recommendations"]
            
            # Enhance language and clarity
            if "special_instructions" in enhanced_data:
                enhanced_instructions = []
                for instruction in enhanced_data["special_instructions"]:
                    # AI-enhanced instruction clarity
                    enhanced_instruction = await self._enhance_instruction_clarity(instruction)
                    enhanced_instructions.append(enhanced_instruction)
                enhanced_data["special_instructions"] = enhanced_instructions
            
            return enhanced_data
            
        except Exception as e:
            logger.warning(f"AI content enhancement failed: {str(e)}")
            return user_data

    async def _enhance_instruction_clarity(self, instruction: str) -> str:
        """Enhance instruction clarity using AI"""
        # Simple enhancement for now - can be expanded with more sophisticated AI
        enhanced = instruction.strip()
        
        # Ensure proper sentence structure
        if not enhanced.endswith('.'):
            enhanced += '.'
        
        # Capitalize first letter
        if enhanced:
            enhanced = enhanced[0].upper() + enhanced[1:]
        
        return enhanced

    def _prepare_will_context(self, user_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare template context for will generation"""
        context = {
            "testator": user_data.get("testator", {}),
            "executors": user_data.get("executors", []),
            "alternate_executors": user_data.get("alternate_executors", []),
            "specific_bequests": user_data.get("specific_bequests", []),
            "residuary_beneficiary": user_data.get("residuary_beneficiary", ""),
            "guardian_appointment": user_data.get("guardian_appointment"),
            "minors": user_data.get("has_minor_children", False),
            "witness1": user_data.get("witnesses", [{}])[0] if user_data.get("witnesses") else {},
            "witness2": user_data.get("witnesses", [{}])[1] if len(user_data.get("witnesses", [])) > 1 else {},
            "current_year": datetime.now().year
        }
        
        # Add pronouns based on gender
        gender = user_data.get("testator", {}).get("gender", "").lower()
        if gender == "male":
            context["testator"]["pronoun_possessive"] = "his"
        elif gender == "female":
            context["testator"]["pronoun_possessive"] = "her"
        else:
            context["testator"]["pronoun_possessive"] = "their"
        
        return context

    def _prepare_poa_context(self, user_data: Dict[str, Any], poa_type: str) -> Dict[str, Any]:
        """Prepare template context for POA generation"""
        context = {
            "grantor": user_data.get("grantor", {}),
            "attorney": user_data.get("attorneys", [{}])[0] if user_data.get("attorneys") else {},
            "alternate_attorney": user_data.get("alternate_attorneys", [{}])[0] if user_data.get("alternate_attorneys") else None,
            "conditions": user_data.get("conditions", []),
            "powers_commence": user_data.get("powers_commence", "immediately"),
            "compensation": user_data.get("attorney_compensation", True),
            "witness": user_data.get("witnesses", [{}])[0] if user_data.get("witnesses") else {},
            "current_year": datetime.now().year
        }
        
        return context

    async def _generate_docx_will(self, content: str, context: Dict[str, Any]) -> Optional[str]:
        """Generate professional DOCX will document"""
        if not DOCX_AVAILABLE:
            return None
            
        try:
            doc = Document()
            
            # Apply Ontario legal styling
            self._apply_ontario_legal_styling(doc)
            
            # Document title
            title = doc.add_heading("LAST WILL AND TESTAMENT", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parse and format content
            sections = content.split('\n\n')
            for section in sections:
                if section.strip():
                    if section.strip().startswith(('1.', '2.', '3.', '4.', '5.')):
                        # Section heading
                        p = doc.add_paragraph()
                        p.add_run(section.strip()).bold = True
                        p.paragraph_format.space_before = Pt(18)
                        p.paragraph_format.space_after = Pt(12)
                    elif 'SIGNED' in section or 'WITNESS' in section:
                        # Signature section
                        p = doc.add_paragraph(section.strip())
                        p.paragraph_format.space_before = Pt(36)
                    else:
                        # Regular paragraph
                        doc.add_paragraph(section.strip())
            
            # Save document
            output_path = f"/tmp/ontario_will_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"DOCX will generation failed: {str(e)}")
            return None

    async def _generate_docx_poa(self, content: str, context: Dict[str, Any], poa_type: str) -> Optional[str]:
        """Generate professional DOCX POA document"""
        if not DOCX_AVAILABLE:
            return None
            
        try:
            doc = Document()
            
            # Apply Ontario legal styling
            self._apply_ontario_legal_styling(doc)
            
            # Document title
            title_text = "POWER OF ATTORNEY FOR PROPERTY" if "property" in poa_type else "POWER OF ATTORNEY FOR PERSONAL CARE"
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parse and format content
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Save document
            output_path = f"/tmp/ontario_poa_{poa_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"DOCX POA generation failed: {str(e)}")
            return None

    def _apply_ontario_legal_styling(self, doc):
        """Apply professional Ontario legal document styling"""
        if not DOCX_AVAILABLE:
            return
            
        try:
            # Set document font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            
            # Set paragraph formatting
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing = 1.15
            paragraph_format.space_after = Pt(6)
            
        except Exception as e:
            logger.warning(f"Styling application failed: {str(e)}")

    async def _generate_pdf_document(self, content: str, document_type: str) -> Optional[str]:
        """Generate professional PDF document"""
        if not REPORTLAB_AVAILABLE:
            return None
            
        try:
            output_path = f"/tmp/ontario_{document_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=inch,
                leftMargin=inch,
                topMargin=inch,
                bottomMargin=inch
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom Ontario legal style
            ontario_style = ParagraphStyle(
                'OntarioLegal',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=12,
                leading=14,
                alignment=0,  # Left align
                spaceAfter=6
            )
            
            title_style = ParagraphStyle(
                'OntarioTitle',
                parent=styles['Title'],
                fontName='Times-Bold',
                fontSize=16,
                alignment=1,  # Center align
                spaceAfter=24
            )
            
            # Build document content
            story = []
            
            # Parse content into paragraphs
            paragraphs = content.split('\n\n')
            
            for i, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    if i == 0 or 'WILL AND TESTAMENT' in paragraph or 'POWER OF ATTORNEY' in paragraph:
                        # Title
                        story.append(Paragraph(paragraph.strip(), title_style))
                    else:
                        # Regular paragraph
                        story.append(Paragraph(paragraph.strip(), ontario_style))
                    
                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            return output_path
            
        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}")
            return None

    def get_document_preview(self, document_type: str, user_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate document preview without full formatting"""
        try:
            if document_type == "will":
                template = self.template_engine.get_template("ontario_will.txt")
                context = self._prepare_will_context(user_data)
            elif document_type in ["poa_property", "poa_personal_care"]:
                template_name = f"ontario_poa_{document_type.split('_')[-1]}.txt"
                template = self.template_engine.get_template(template_name)
                context = self._prepare_poa_context(user_data, document_type)
            else:
                return {"success": False, "error": "Unsupported document type"}
            
            content = template.render(**context)
            
            return {
                "success": True,
                "preview": content,
                "word_count": len(content.split()),
                "estimated_pages": max(1, len(content) // 2500)  # Rough estimate
            }
            
        except Exception as e:
            logger.error(f"Document preview generation failed: {str(e)}")
            return {"success": False, "error": str(e)}

    def validate_document_data(self, document_type: str, user_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate document data completeness and correctness"""
        validation_errors = []
        warnings = []
        
        if document_type == "will":
            # Required fields for will
            required_fields = ["testator", "executors", "residuary_beneficiary", "witnesses"]
            
            for field in required_fields:
                if not user_data.get(field):
                    validation_errors.append(f"Missing required field: {field}")
            
            # Validate executors
            executors = user_data.get("executors", [])
            if len(executors) < 1:
                validation_errors.append("At least one executor is required")
            
            # Validate witnesses
            witnesses = user_data.get("witnesses", [])
            if len(witnesses) < 2:
                validation_errors.append("Two witnesses are required for a valid will")
                
        elif document_type in ["poa_property", "poa_personal_care"]:
            # Required fields for POA
            required_fields = ["grantor", "attorneys", "witnesses"]
            
            for field in required_fields:
                if not user_data.get(field):
                    validation_errors.append(f"Missing required field: {field}")
            
            # Validate attorneys
            attorneys = user_data.get("attorneys", [])
            if len(attorneys) < 1:
                validation_errors.append("At least one attorney is required")
        
        return {
            "valid": len(validation_errors) == 0,
            "errors": validation_errors,
            "warnings": warnings
        }

    def get_supported_document_types(self) -> List[str]:
        """Get list of supported document types"""
        return ["will", "poa_property", "poa_personal_care"]

    def get_template_fields(self, document_type: str) -> Dict[str, Any]:
        """Get required and optional fields for document type"""
        field_definitions = {
            "will": {
                "required": [
                    "testator.full_name",
                    "testator.address", 
                    "testator.city",
                    "testator.postal_code",
                    "executors",
                    "residuary_beneficiary",
                    "witnesses"
                ],
                "optional": [
                    "alternate_executors",
                    "specific_bequests",
                    "guardian_appointment",
                    "special_instructions"
                ]
            },
            "poa_property": {
                "required": [
                    "grantor.full_name",
                    "grantor.address",
                    "grantor.city", 
                    "grantor.postal_code",
                    "attorneys",
                    "witnesses"
                ],
                "optional": [
                    "alternate_attorneys",
                    "conditions",
                    "powers_commence",
                    "compensation"
                ]
            },
            "poa_personal_care": {
                "required": [
                    "grantor.full_name",
                    "grantor.address",
                    "grantor.city",
                    "grantor.postal_code", 
                    "attorneys",
                    "witnesses"
                ],
                "optional": [
                    "alternate_attorneys",
                    "healthcare_instructions",
                    "end_of_life_wishes",
                    "conditions"
                ]
            }
        }
        
        return field_definitions.get(document_type, {"required": [], "optional": []})

    def is_ready(self) -> bool:
        """Check if document generator is ready"""
        return self.is_initialized

    async def health_check(self) -> Dict[str, Any]:
        """Health check for document generator"""
        return {
            "status": "ready" if self.is_initialized else "not_ready",
            "capabilities": {
                "docx_generation": DOCX_AVAILABLE,
                "pdf_generation": REPORTLAB_AVAILABLE or WEASYPRINT_AVAILABLE,
                "ai_enhancement": self.ai_service.is_ready() if self.ai_service else False,
                "template_engine": self.template_engine is not None
            },
            "supported_documents": self.get_supported_document_types(),
            "timestamp": datetime.now().isoformat()
        }

# Import asdict for analysis conversion
try:
    from dataclasses import asdict
except ImportError:
    def asdict(obj):
        """Fallback asdict implementation"""
        if hasattr(obj, '__dict__'):
            return obj.__dict__
        return {}