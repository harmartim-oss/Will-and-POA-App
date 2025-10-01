# backend/core/document_combiner.py
"""
Document Combiner - Merge multiple documents with custom page numbering
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
import io
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)


class DocumentCombiner:
    """Combine multiple documents into a single document with page numbering"""
    
    @staticmethod
    def add_page_number_to_pdf(pdf_bytes: bytes, page_num: int, start_page: int = 2) -> bytes:
        """Add page number to PDF (only starting from specified page)"""
        try:
            # Create a PDF reader
            existing_pdf = PdfReader(io.BytesIO(pdf_bytes))
            output = PdfWriter()
            
            for i, page in enumerate(existing_pdf.pages):
                # Create page number overlay only if page >= start_page
                if i + 1 >= start_page:
                    packet = io.BytesIO()
                    can = canvas.Canvas(packet, pagesize=letter)
                    
                    # Add page number in top right (page number starts from 2 for second page)
                    display_page_num = page_num + (i - start_page + 2)
                    can.setFont("Helvetica", 10)
                    can.drawRightString(7.5*inch, 10.5*inch, str(display_page_num))
                    
                    can.save()
                    packet.seek(0)
                    
                    # Merge the page number with the existing page
                    overlay_pdf = PdfReader(packet)
                    page.merge_page(overlay_pdf.pages[0])
                
                output.add_page(page)
            
            # Write to bytes
            output_stream = io.BytesIO()
            output.write(output_stream)
            output_stream.seek(0)
            return output_stream.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to add page numbers: {str(e)}")
            return pdf_bytes  # Return original if numbering fails
    
    @staticmethod
    async def combine_pdfs(
        documents: List[Dict[str, Any]],
        add_page_numbers: bool = True,
        start_numbering_page: int = 2
    ) -> bytes:
        """
        Combine multiple PDF documents into one
        
        Args:
            documents: List of dicts with 'content' (bytes) and 'name' (str)
            add_page_numbers: Whether to add page numbers
            start_numbering_page: Which page to start numbering (1-indexed, default 2)
        
        Returns:
            Combined PDF as bytes
        """
        try:
            merger = PdfMerger()
            current_page = 0
            
            if add_page_numbers:
                # Process each document with page numbers
                processed_docs = []
                
                for doc in documents:
                    doc_bytes = doc['content']
                    
                    # Add page numbers to this document
                    numbered_doc = DocumentCombiner.add_page_number_to_pdf(
                        doc_bytes,
                        current_page,
                        start_numbering_page
                    )
                    
                    processed_docs.append(numbered_doc)
                    
                    # Count pages for next document
                    reader = PdfReader(io.BytesIO(doc_bytes))
                    current_page += len(reader.pages)
                
                # Merge all processed documents
                for doc_bytes in processed_docs:
                    merger.append(io.BytesIO(doc_bytes))
            else:
                # Simple merge without page numbers
                for doc in documents:
                    merger.append(io.BytesIO(doc['content']))
            
            # Write merged PDF
            output = io.BytesIO()
            merger.write(output)
            merger.close()
            output.seek(0)
            
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to combine PDFs: {str(e)}")
            raise
    
    @staticmethod
    async def combine_docx(
        documents: List[Dict[str, Any]],
        add_page_numbers: bool = True,
        start_numbering_page: int = 2
    ) -> bytes:
        """
        Combine multiple DOCX documents into one
        
        Args:
            documents: List of dicts with 'content' (bytes) and 'name' (str)
            add_page_numbers: Whether to add page numbers
            start_numbering_page: Which page to start numbering (1-indexed, default 2)
        
        Returns:
            Combined DOCX as bytes
        """
        try:
            # Create main document from first document
            main_doc = Document(io.BytesIO(documents[0]['content']))
            
            # Add remaining documents
            for i, doc_info in enumerate(documents[1:], 1):
                # Add page break before next document
                main_doc.add_page_break()
                
                # Load next document
                next_doc = Document(io.BytesIO(doc_info['content']))
                
                # Copy all elements from next document
                for element in next_doc.element.body:
                    main_doc.element.body.append(element)
            
            # Add page numbers if requested
            if add_page_numbers:
                DocumentCombiner._add_page_numbers_to_docx(main_doc, start_numbering_page)
            
            # Save to bytes
            buffer = io.BytesIO()
            main_doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to combine DOCX files: {str(e)}")
            raise
    
    @staticmethod
    def _add_page_numbers_to_docx(doc: Document, start_page: int = 2):
        """Add page numbers to DOCX document in top right corner"""
        try:
            # Create page numbering in header
            for section in doc.sections:
                header = section.header
                
                # Clear existing header content
                for paragraph in header.paragraphs:
                    paragraph.clear()
                
                # Add page number field
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Add field code for page number
                run = paragraph.add_run()
                
                # Create field element
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = 'PAGE'
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
                
                # Set font size
                run.font.size = Pt(10)
                
                # Different first page if start_page > 1
                if start_page > 1:
                    section.different_first_page_header_footer = True
                    # First page header should be empty (no page number)
                    first_page_header = section.first_page_header
                    for paragraph in first_page_header.paragraphs:
                        paragraph.clear()
                        
        except Exception as e:
            logger.error(f"Failed to add page numbers to DOCX: {str(e)}")


class BatchDocumentGenerator:
    """Generate and package multiple documents for batch download"""
    
    @staticmethod
    async def create_batch_download(
        documents: List[Dict[str, Any]],
        format: str = "zip"
    ) -> bytes:
        """
        Create a batch download package
        
        Args:
            documents: List of dicts with 'content' (bytes), 'name' (str), 'type' (str)
            format: 'zip' or 'combined' (single document)
        
        Returns:
            Package as bytes
        """
        if format == "combined":
            # Determine document type from first document
            doc_type = documents[0].get('type', 'pdf')
            
            if doc_type == 'pdf':
                return await DocumentCombiner.combine_pdfs(documents, add_page_numbers=True)
            elif doc_type == 'docx':
                return await DocumentCombiner.combine_docx(documents, add_page_numbers=True)
            else:
                raise ValueError(f"Unsupported document type for combining: {doc_type}")
        else:
            # Create ZIP archive
            import zipfile
            
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for doc in documents:
                    # Add each document to zip
                    filename = doc.get('name', f"document_{documents.index(doc) + 1}")
                    # Ensure proper extension
                    if not any(filename.endswith(ext) for ext in ['.pdf', '.docx', '.doc', '.txt']):
                        doc_type = doc.get('type', 'pdf')
                        filename = f"{filename}.{doc_type}"
                    
                    zip_file.writestr(filename, doc['content'])
            
            zip_buffer.seek(0)
            return zip_buffer.getvalue()
