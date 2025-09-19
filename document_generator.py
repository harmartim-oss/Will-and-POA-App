import os
import tempfile
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from docx import Document as WordDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for legal documents"""
        # Title style
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Times-Bold'
        )
        
        # Heading style
        self.heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            fontName='Times-Bold'
        )
        
        # Body style
        self.body_style = ParagraphStyle(
            'CustomBody',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            fontName='Times-Roman',
            leftIndent=0.5*inch,
            rightIndent=0.5*inch
        )
        
        # Signature style
        self.signature_style = ParagraphStyle(
            'Signature',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=24,
            spaceBefore=24,
            fontName='Times-Roman'
        )
    
    def generate_pdf(self, document):
        """Generate PDF document"""
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        doc = SimpleDocTemplate(
            temp_file.name,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        story = []
        content = document.get_content()
        
        if document.document_type == 'will':
            story = self._generate_will_content(content)
        elif document.document_type == 'poa_property':
            story = self._generate_poa_property_content(content)
        elif document.document_type == 'poa_care':
            story = self._generate_poa_care_content(content)
        
        doc.build(story)
        temp_file.close()
        
        return temp_file.name
    
    def generate_word(self, document):
        """Generate Word document"""
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc = WordDocument()
        
        content = document.get_content()
        
        if document.document_type == 'will':
            self._generate_will_word(doc, content)
        elif document.document_type == 'poa_property':
            self._generate_poa_property_word(doc, content)
        elif document.document_type == 'poa_care':
            self._generate_poa_care_word(doc, content)
        
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name
    
    def _generate_will_content(self, content):
        """Generate will content for PDF"""
        story = []
        
        # Title
        story.append(Paragraph("LAST WILL AND TESTAMENT", self.title_style))
        story.append(Spacer(1, 12))
        
        # Personal information
        personal_info = content.get('personal_info', {})
        story.append(Paragraph(f"OF {personal_info.get('full_name', '[NAME]').upper()}", self.title_style))
        story.append(Spacer(1, 24))
        
        # Declaration
        story.append(Paragraph("DECLARATION", self.heading_style))
        declaration_text = f"""I, {personal_info.get('full_name', '[NAME]')}, of {personal_info.get('address', '[ADDRESS]')}, 
        in the Province of Ontario, being of sound mind and disposing memory and not acting under duress, menace, fraud, 
        or undue influence of any person whomsoever, do make, publish and declare this to be my Last Will and Testament, 
        hereby revoking all former Wills and Codicils by me at any time heretofore made."""
        story.append(Paragraph(declaration_text, self.body_style))
        story.append(Spacer(1, 12))
        
        # Executor appointment
        executor_info = content.get('executor', {})
        story.append(Paragraph("APPOINTMENT OF EXECUTOR", self.heading_style))
        executor_text = f"""I appoint {executor_info.get('name', '[EXECUTOR NAME]')} of {executor_info.get('address', '[EXECUTOR ADDRESS]')} 
        to be the Executor of this my Will."""
        story.append(Paragraph(executor_text, self.body_style))
        story.append(Spacer(1, 12))
        
        # Bequests
        bequests = content.get('bequests', [])
        if bequests:
            story.append(Paragraph("SPECIFIC BEQUESTS", self.heading_style))
            for i, bequest in enumerate(bequests, 1):
                bequest_text = f"{i}. I give and bequeath {bequest.get('item', '[ITEM]')} to {bequest.get('beneficiary', '[BENEFICIARY]')}."
                story.append(Paragraph(bequest_text, self.body_style))
            story.append(Spacer(1, 12))
        
        # Residuary estate
        residuary = content.get('residuary', {})
        story.append(Paragraph("RESIDUARY ESTATE", self.heading_style))
        residuary_text = f"""I give, devise and bequeath all the rest, residue and remainder of my estate, 
        both real and personal, of whatsoever nature and wheresoever situate, to {residuary.get('beneficiary', '[RESIDUARY BENEFICIARY]')}."""
        story.append(Paragraph(residuary_text, self.body_style))
        story.append(Spacer(1, 24))
        
        # Signature section
        story.append(Paragraph("IN WITNESS WHEREOF I have hereunto set my hand this _____ day of _____________, 20____.", self.signature_style))
        story.append(Spacer(1, 24))
        story.append(Paragraph("_________________________________", self.signature_style))
        story.append(Paragraph(f"{personal_info.get('full_name', '[NAME]')}, Testator", self.signature_style))
        story.append(Spacer(1, 24))
        
        # Witness section
        story.append(Paragraph("SIGNED, PUBLISHED AND DECLARED by the above-named Testator as and for his/her Last Will and Testament, in the presence of us, both present at the same time, who at his/her request, in his/her presence, and in the presence of each other, have hereunto subscribed our names as witnesses.", self.body_style))
        story.append(Spacer(1, 24))
        
        story.append(Paragraph("_________________________________", self.signature_style))
        story.append(Paragraph("Witness #1 Signature", self.signature_style))
        story.append(Paragraph("_________________________________", self.signature_style))
        story.append(Paragraph("Witness #1 Name (Print)", self.signature_style))
        story.append(Paragraph("_________________________________", self.signature_style))
        story.append(Paragraph("Witness #1 Address", self.signature_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("_________________________________", self.signature_style))
        story.append(Paragraph("Witness #2 Signature", self.signature_style))
        story.append(Paragraph("_________________________________", self.signature_style))
        story.append(Paragraph("Witness #2 Name (Print)", self.signature_style))
        story.append(Paragraph("_________________________________", self.signature_style))
        story.append(Paragraph("Witness #2 Address", self.signature_style))
        
        return story
    
    def _generate_poa_property_content(self, content):
        """Generate power of attorney for property content for PDF"""
        story = []
        
        # Title
        story.append(Paragraph("CONTINUING POWER OF ATTORNEY FOR PROPERTY", self.title_style))
        story.append(Spacer(1, 24))
        
        # Grantor information
        grantor_info = content.get('grantor_info', {})
        story.append(Paragraph("GRANTOR INFORMATION", self.heading_style))
        grantor_text = f"""I, {grantor_info.get('full_name', '[NAME]')}, of {grantor_info.get('address', '[ADDRESS]')}, 
        in the Province of Ontario, being of sound mind, do hereby appoint the person(s) named below as my attorney(s) 
        for property with authority to act on my behalf in accordance with the Substitute Decisions Act, 1992."""
        story.append(Paragraph(grantor_text, self.body_style))
        story.append(Spacer(1, 12))
        
        # Attorney appointment
        attorney_info = content.get('attorney', {})
        story.append(Paragraph("APPOINTMENT OF ATTORNEY", self.heading_style))
        attorney_text = f"""I appoint {attorney_info.get('name', '[ATTORNEY NAME]')} of {attorney_info.get('address', '[ATTORNEY ADDRESS]')} 
        to be my attorney for property."""
        story.append(Paragraph(attorney_text, self.body_style))
        story.append(Spacer(1, 12))
        
        # Powers granted
        powers = content.get('powers', [])
        if powers:
            story.append(Paragraph("POWERS GRANTED", self.heading_style))
            story.append(Paragraph("I grant my attorney the following powers:", self.body_style))
            for power in powers:
                story.append(Paragraph(f"• {power}", self.body_style))
            story.append(Spacer(1, 12))
        
        # Conditions and restrictions
        conditions = content.get('conditions', '')
        if conditions:
            story.append(Paragraph("CONDITIONS AND RESTRICTIONS", self.heading_style))
            story.append(Paragraph(conditions, self.body_style))
            story.append(Spacer(1, 12))
        
        # Signature section
        story.append(Spacer(1, 24))
        story.append(Paragraph("IN WITNESS WHEREOF I have executed this Power of Attorney this _____ day of _____________, 20____.", self.signature_style))
        story.append(Spacer(1, 24))
        story.append(Paragraph("_________________________________", self.signature_style))
        story.append(Paragraph(f"{grantor_info.get('full_name', '[NAME]')}, Grantor", self.signature_style))
        
        return story
    
    def _generate_poa_care_content(self, content):
        """Generate power of attorney for personal care content for PDF"""
        story = []
        
        # Title
        story.append(Paragraph("POWER OF ATTORNEY FOR PERSONAL CARE", self.title_style))
        story.append(Spacer(1, 24))
        
        # Similar structure to property POA but for personal care decisions
        # Implementation would follow similar pattern
        
        return story
    
    def _generate_will_word(self, doc, content):
        """Generate will content for Word document"""
        # Add title
        title = doc.add_heading('LAST WILL AND TESTAMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add personal information
        personal_info = content.get('personal_info', {})
        name_para = doc.add_heading(f"OF {personal_info.get('full_name', '[NAME]').upper()}", level=1)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add declaration
        doc.add_heading('DECLARATION', level=2)
        declaration_text = f"""I, {personal_info.get('full_name', '[NAME]')}, of {personal_info.get('address', '[ADDRESS]')}, 
        in the Province of Ontario, being of sound mind and disposing memory and not acting under duress, menace, fraud, 
        or undue influence of any person whomsoever, do make, publish and declare this to be my Last Will and Testament, 
        hereby revoking all former Wills and Codicils by me at any time heretofore made."""
        doc.add_paragraph(declaration_text)
        
        # Continue with similar structure as PDF version
        # Implementation would follow similar pattern for all sections
    
    def _generate_poa_property_word(self, doc, content):
        """Generate power of attorney for property content for Word document"""
        # Similar implementation for Word format
        pass
    
    def _generate_poa_care_word(self, doc, content):
        """Generate power of attorney for personal care content for Word document"""
        # Similar implementation for Word format
        pass

