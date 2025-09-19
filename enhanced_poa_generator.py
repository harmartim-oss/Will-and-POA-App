"""
Enhanced Power of Attorney Generator for Ontario
Comprehensive implementation for both Property and Personal Care POAs
"""

import json
import logging
from datetime import datetime, date
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, asdict
from enum import Enum
import re
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class POAType(Enum):
    """Types of Power of Attorney"""
    PROPERTY = "property"
    PERSONAL_CARE = "personal_care"
    COMBINED = "combined"

class POAPropertyType(Enum):
    """Types of Property POA"""
    CONTINUING = "continuing"
    NON_CONTINUING = "non_continuing"

@dataclass
class PersonInfo:
    """Information about a person (grantor, attorney, witness)"""
    full_name: str
    address: str
    city: str
    province: str
    postal_code: str
    phone: Optional[str] = None
    email: Optional[str] = None
    relationship: Optional[str] = None
    date_of_birth: Optional[str] = None

@dataclass
class PropertyPowers:
    """Specific powers for Property POA"""
    real_estate: bool = True
    personal_property: bool = True
    banking: bool = True
    investments: bool = True
    insurance: bool = True
    business_interests: bool = True
    legal_proceedings: bool = True
    tax_matters: bool = True
    government_benefits: bool = True
    employment_matters: bool = False
    custom_powers: List[str] = None

@dataclass
class PersonalCarePowers:
    """Specific powers for Personal Care POA"""
    health_care: bool = True
    nutrition: bool = True
    shelter: bool = True
    clothing: bool = True
    hygiene: bool = True
    safety: bool = True
    social_activities: bool = True
    employment: bool = False
    custom_instructions: List[str] = None
    medical_preferences: Dict[str, Any] = None

@dataclass
class POARestrictions:
    """Restrictions and limitations on POA"""
    cannot_make_will: bool = True
    cannot_change_beneficiaries: bool = True
    cannot_delegate_authority: bool = False
    financial_limits: Optional[Dict[str, float]] = None
    specific_restrictions: List[str] = None
    requires_consent_for: List[str] = None

@dataclass
class POADocument:
    """Complete Power of Attorney document data"""
    poa_type: POAType
    property_type: Optional[POAPropertyType] = None
    grantor: PersonInfo = None
    attorneys: List[PersonInfo] = None
    substitute_attorneys: List[PersonInfo] = None
    witnesses: List[PersonInfo] = None
    property_powers: Optional[PropertyPowers] = None
    personal_care_powers: Optional[PersonalCarePowers] = None
    restrictions: Optional[POARestrictions] = None
    special_instructions: List[str] = None
    execution_date: Optional[str] = None
    effective_date: Optional[str] = None
    revocation_conditions: List[str] = None
    compensation: Optional[Dict[str, Any]] = None
    bond_required: bool = False
    accounting_requirements: Optional[Dict[str, Any]] = None

class EnhancedPOAGenerator:
    """
    Enhanced Power of Attorney generator with comprehensive Ontario compliance
    """
    
    def __init__(self):
        self.ontario_requirements = self._load_ontario_requirements()
        self.legal_templates = self._load_legal_templates()
        self.styles = self._initialize_styles()
    
    def _load_ontario_requirements(self) -> Dict[str, Any]:
        """Load detailed Ontario POA requirements"""
        return {
            "property_poa": {
                "minimum_age": 18,
                "capacity_requirements": [
                    "Understand the nature and effect of the document",
                    "Understand the extent of property covered",
                    "Know that the attorney will be able to do anything with the property that the grantor could do",
                    "Know that the attorney must account for dealings with the property",
                    "Know that the grantor may revoke the POA if capable",
                    "Appreciate that unless the POA provides otherwise, the attorney will not be able to benefit from the property"
                ],
                "witness_requirements": {
                    "minimum_count": 2,
                    "disqualifications": [
                        "The attorney or the attorney's spouse or partner",
                        "A child of the grantor, if the attorney is the grantor's spouse or partner",
                        "A person who is not the grantor but who signed the POA on behalf of the grantor",
                        "A person who is less than 18 years old"
                    ]
                },
                "mandatory_clauses": [
                    "Appointment of attorney",
                    "Description of powers granted",
                    "Continuing or non-continuing designation",
                    "Grantor's signature",
                    "Witness signatures and information"
                ]
            },
            "personal_care_poa": {
                "minimum_age": 16,
                "capacity_requirements": [
                    "Understand the nature and effect of the document",
                    "Understand what types of personal care may be needed",
                    "Know that the attorney will be making personal care decisions",
                    "Appreciate the consequences of giving the POA"
                ],
                "witness_requirements": {
                    "minimum_count": 2,
                    "disqualifications": [
                        "The attorney or the attorney's spouse or partner",
                        "A person who is providing health care, residential, social, training or support services to the grantor for compensation",
                        "A person who is less than 18 years old"
                    ]
                },
                "scope_limitations": [
                    "Cannot consent to treatment if grantor is capable of consenting",
                    "Cannot override grantor's prior capable wishes",
                    "Must act in grantor's best interests",
                    "Cannot consent to experimental treatment without court approval"
                ]
            }
        }
    
    def _load_legal_templates(self) -> Dict[str, str]:
        """Load legal clause templates"""
        return {
            "property_appointment": """I, {grantor_name}, of {grantor_address}, being of sound mind and understanding the nature and effect of this document, DO HEREBY APPOINT {attorney_names} of {attorney_addresses}, jointly and severally, to be my attorney(s) for property.""",
            
            "continuing_clause": """This Power of Attorney for Property SHALL CONTINUE to be effective if I become mentally incapable of managing my property, subject to the Substitute Decisions Act, 1992.""",
            
            "non_continuing_clause": """This Power of Attorney for Property SHALL NOT continue to be effective if I become mentally incapable of managing my property.""",
            
            "property_powers_general": """My attorney(s) may do on my behalf anything in respect of property that I could do if capable of managing property, except make a will, subject to the law and to any conditions or restrictions contained in this document.""",
            
            "personal_care_appointment": """I, {grantor_name}, of {grantor_address}, being at least 16 years of age and mentally capable of giving this Power of Attorney for Personal Care, DO HEREBY APPOINT {attorney_names} of {attorney_addresses}, to be my attorney(s) for personal care.""",
            
            "personal_care_powers": """My attorney(s) may make decisions concerning my personal care including decisions about my health care, nutrition, shelter, clothing, hygiene and safety.""",
            
            "revocation_clause": """I HEREBY REVOKE any prior Power of Attorney for {poa_type} made by me.""",
            
            "witness_attestation": """SIGNED, SEALED AND DELIVERED in the presence of us, who in the presence of the grantor and of each other, have signed our names as witnesses.""",
            
            "capacity_statement": """I understand the nature and effect of this Power of Attorney, and I am signing it voluntarily."""
        }
    
    def _initialize_styles(self) -> Dict[str, Any]:
        """Initialize document styles"""
        styles = getSampleStyleSheet()
        
        # Custom styles for legal documents
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=16,
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        )
        
        signature_style = ParagraphStyle(
            'SignatureStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=24,
            spaceBefore=12,
            fontName='Helvetica'
        )
        
        return {
            'title': title_style,
            'heading': heading_style,
            'body': body_style,
            'signature': signature_style
        }
    
    def generate_property_poa(self, poa_data: POADocument) -> Dict[str, Any]:
        """Generate a comprehensive Property Power of Attorney"""
        try:
            # Validate data
            validation_result = self._validate_property_poa_data(poa_data)
            if not validation_result["valid"]:
                return {"success": False, "errors": validation_result["errors"]}
            
            # Generate document content
            content = self._build_property_poa_content(poa_data)
            
            # Create PDF and Word versions
            pdf_path = f"/tmp/property_poa_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            word_path = f"/tmp/property_poa_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            pdf_success = self._create_pdf_document(content, pdf_path, "POWER OF ATTORNEY FOR PROPERTY")
            word_success = self._create_word_document(content, word_path, "POWER OF ATTORNEY FOR PROPERTY")
            
            return {
                "success": True,
                "pdf_path": pdf_path if pdf_success else None,
                "word_path": word_path if word_success else None,
                "content": content,
                "validation": validation_result
            }
            
        except Exception as e:
            logger.error(f"Error generating Property POA: {e}")
            return {"success": False, "error": str(e)}
    
    def generate_personal_care_poa(self, poa_data: POADocument) -> Dict[str, Any]:
        """Generate a comprehensive Personal Care Power of Attorney"""
        try:
            # Validate data
            validation_result = self._validate_personal_care_poa_data(poa_data)
            if not validation_result["valid"]:
                return {"success": False, "errors": validation_result["errors"]}
            
            # Generate document content
            content = self._build_personal_care_poa_content(poa_data)
            
            # Create PDF and Word versions
            pdf_path = f"/tmp/personal_care_poa_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            word_path = f"/tmp/personal_care_poa_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            pdf_success = self._create_pdf_document(content, pdf_path, "POWER OF ATTORNEY FOR PERSONAL CARE")
            word_success = self._create_word_document(content, word_path, "POWER OF ATTORNEY FOR PERSONAL CARE")
            
            return {
                "success": True,
                "pdf_path": pdf_path if pdf_success else None,
                "word_path": word_path if word_success else None,
                "content": content,
                "validation": validation_result
            }
            
        except Exception as e:
            logger.error(f"Error generating Personal Care POA: {e}")
            return {"success": False, "error": str(e)}
    
    def generate_combined_poa(self, poa_data: POADocument) -> Dict[str, Any]:
        """Generate combined Property and Personal Care POA documents"""
        try:
            # Generate both documents
            property_result = self.generate_property_poa(poa_data)
            personal_care_result = self.generate_personal_care_poa(poa_data)
            
            return {
                "success": property_result["success"] and personal_care_result["success"],
                "property_poa": property_result,
                "personal_care_poa": personal_care_result
            }
            
        except Exception as e:
            logger.error(f"Error generating combined POA: {e}")
            return {"success": False, "error": str(e)}
    
    def _validate_property_poa_data(self, poa_data: POADocument) -> Dict[str, Any]:
        """Validate Property POA data against Ontario requirements"""
        errors = []
        warnings = []
        
        # Check grantor information
        if not poa_data.grantor:
            errors.append("Grantor information is required")
        else:
            if not poa_data.grantor.full_name:
                errors.append("Grantor full name is required")
            if not poa_data.grantor.address:
                errors.append("Grantor address is required")
        
        # Check attorney information
        if not poa_data.attorneys or len(poa_data.attorneys) == 0:
            errors.append("At least one attorney must be appointed")
        else:
            for i, attorney in enumerate(poa_data.attorneys):
                if not attorney.full_name:
                    errors.append(f"Attorney {i+1} full name is required")
                if not attorney.address:
                    errors.append(f"Attorney {i+1} address is required")
        
        # Check witness requirements
        if not poa_data.witnesses or len(poa_data.witnesses) < 2:
            errors.append("At least 2 witnesses are required for Property POA")
        else:
            for i, witness in enumerate(poa_data.witnesses):
                if not witness.full_name:
                    errors.append(f"Witness {i+1} full name is required")
                if not witness.address:
                    errors.append(f"Witness {i+1} address is required")
        
        # Check for witness disqualifications
        if poa_data.witnesses and poa_data.attorneys:
            attorney_names = [a.full_name.lower() for a in poa_data.attorneys]
            for witness in poa_data.witnesses:
                if witness.full_name.lower() in attorney_names:
                    errors.append("A witness cannot be the same person as an attorney")
        
        # Check property type specification
        if not poa_data.property_type:
            warnings.append("Property POA type (continuing/non-continuing) should be specified")
        
        # Check execution date
        if not poa_data.execution_date:
            warnings.append("Execution date should be specified")
        
        return {
            "valid": len(errors) == 0,
            "errors": errors,
            "warnings": warnings
        }
    
    def _validate_personal_care_poa_data(self, poa_data: POADocument) -> Dict[str, Any]:
        """Validate Personal Care POA data against Ontario requirements"""
        errors = []
        warnings = []
        
        # Check grantor information
        if not poa_data.grantor:
            errors.append("Grantor information is required")
        else:
            if not poa_data.grantor.full_name:
                errors.append("Grantor full name is required")
            if not poa_data.grantor.address:
                errors.append("Grantor address is required")
        
        # Check attorney information
        if not poa_data.attorneys or len(poa_data.attorneys) == 0:
            errors.append("At least one attorney must be appointed")
        else:
            for i, attorney in enumerate(poa_data.attorneys):
                if not attorney.full_name:
                    errors.append(f"Attorney {i+1} full name is required")
                if not attorney.address:
                    errors.append(f"Attorney {i+1} address is required")
        
        # Check witness requirements
        if not poa_data.witnesses or len(poa_data.witnesses) < 2:
            errors.append("At least 2 witnesses are required for Personal Care POA")
        else:
            for i, witness in enumerate(poa_data.witnesses):
                if not witness.full_name:
                    errors.append(f"Witness {i+1} full name is required")
                if not witness.address:
                    errors.append(f"Witness {i+1} address is required")
        
        # Check for witness disqualifications
        if poa_data.witnesses and poa_data.attorneys:
            attorney_names = [a.full_name.lower() for a in poa_data.attorneys]
            for witness in poa_data.witnesses:
                if witness.full_name.lower() in attorney_names:
                    errors.append("A witness cannot be the same person as an attorney")
        
        # Check personal care powers
        if not poa_data.personal_care_powers:
            warnings.append("Personal care powers should be specified")
        
        # Check execution date
        if not poa_data.execution_date:
            warnings.append("Execution date should be specified")
        
        return {
            "valid": len(errors) == 0,
            "errors": errors,
            "warnings": warnings
        }
    
    def _build_property_poa_content(self, poa_data: POADocument) -> List[Dict[str, Any]]:
        """Build content structure for Property POA"""
        content = []
        
        # Title
        content.append({
            "type": "title",
            "text": "POWER OF ATTORNEY FOR PROPERTY"
        })
        
        # Grantor identification and appointment
        attorney_names = ", ".join([a.full_name for a in poa_data.attorneys])
        attorney_addresses = "; ".join([f"{a.address}, {a.city}, {a.province} {a.postal_code}" for a in poa_data.attorneys])
        
        appointment_text = self.legal_templates["property_appointment"].format(
            grantor_name=poa_data.grantor.full_name,
            grantor_address=f"{poa_data.grantor.address}, {poa_data.grantor.city}, {poa_data.grantor.province} {poa_data.grantor.postal_code}",
            attorney_names=attorney_names,
            attorney_addresses=attorney_addresses
        )
        
        content.append({
            "type": "paragraph",
            "text": appointment_text
        })
        
        # Continuing/Non-continuing clause
        if poa_data.property_type == POAPropertyType.CONTINUING:
            content.append({
                "type": "paragraph",
                "text": self.legal_templates["continuing_clause"]
            })
        else:
            content.append({
                "type": "paragraph",
                "text": self.legal_templates["non_continuing_clause"]
            })
        
        # Powers granted
        content.append({
            "type": "heading",
            "text": "POWERS GRANTED"
        })
        
        content.append({
            "type": "paragraph",
            "text": self.legal_templates["property_powers_general"]
        })
        
        # Specific powers
        if poa_data.property_powers:
            powers_list = []
            if poa_data.property_powers.real_estate:
                powers_list.append("Buy, sell, mortgage, lease, and otherwise deal with real estate")
            if poa_data.property_powers.banking:
                powers_list.append("Operate bank accounts, make deposits and withdrawals")
            if poa_data.property_powers.investments:
                powers_list.append("Buy, sell, and manage investments and securities")
            if poa_data.property_powers.insurance:
                powers_list.append("Purchase, maintain, and claim on insurance policies")
            if poa_data.property_powers.business_interests:
                powers_list.append("Manage business interests and partnerships")
            if poa_data.property_powers.legal_proceedings:
                powers_list.append("Commence, defend, and settle legal proceedings")
            if poa_data.property_powers.tax_matters:
                powers_list.append("File tax returns and deal with tax authorities")
            if poa_data.property_powers.government_benefits:
                powers_list.append("Apply for and receive government benefits")
            
            if powers_list:
                content.append({
                    "type": "paragraph",
                    "text": "Without limiting the generality of the foregoing, my attorney(s) may:"
                })
                
                for power in powers_list:
                    content.append({
                        "type": "bullet",
                        "text": power
                    })
        
        # Custom powers
        if poa_data.property_powers and poa_data.property_powers.custom_powers:
            content.append({
                "type": "heading",
                "text": "ADDITIONAL POWERS"
            })
            
            for custom_power in poa_data.property_powers.custom_powers:
                content.append({
                    "type": "bullet",
                    "text": custom_power
                })
        
        # Restrictions
        if poa_data.restrictions:
            content.append({
                "type": "heading",
                "text": "RESTRICTIONS AND LIMITATIONS"
            })
            
            if poa_data.restrictions.cannot_make_will:
                content.append({
                    "type": "paragraph",
                    "text": "My attorney(s) cannot make a will on my behalf."
                })
            
            if poa_data.restrictions.financial_limits:
                for limit_type, amount in poa_data.restrictions.financial_limits.items():
                    content.append({
                        "type": "paragraph",
                        "text": f"My attorney(s) cannot {limit_type} more than ${amount:,.2f} without court approval."
                    })
            
            if poa_data.restrictions.specific_restrictions:
                for restriction in poa_data.restrictions.specific_restrictions:
                    content.append({
                        "type": "bullet",
                        "text": restriction
                    })
        
        # Compensation
        if poa_data.compensation:
            content.append({
                "type": "heading",
                "text": "COMPENSATION"
            })
            
            if poa_data.compensation.get("allowed", False):
                if poa_data.compensation.get("amount"):
                    content.append({
                        "type": "paragraph",
                        "text": f"My attorney(s) shall be entitled to compensation of ${poa_data.compensation['amount']:,.2f} per year."
                    })
                else:
                    content.append({
                        "type": "paragraph",
                        "text": "My attorney(s) shall be entitled to reasonable compensation as provided by law."
                    })
            else:
                content.append({
                    "type": "paragraph",
                    "text": "My attorney(s) shall serve without compensation."
                })
        
        # Accounting requirements
        if poa_data.accounting_requirements:
            content.append({
                "type": "heading",
                "text": "ACCOUNTING"
            })
            
            frequency = poa_data.accounting_requirements.get("frequency", "annually")
            content.append({
                "type": "paragraph",
                "text": f"My attorney(s) shall provide an accounting of all transactions {frequency}."
            })
        
        # Revocation clause
        content.append({
            "type": "paragraph",
            "text": self.legal_templates["revocation_clause"].format(poa_type="Property")
        })
        
        # Capacity statement
        content.append({
            "type": "paragraph",
            "text": self.legal_templates["capacity_statement"]
        })
        
        # Execution section
        content.append({
            "type": "heading",
            "text": "EXECUTION"
        })
        
        execution_date = poa_data.execution_date or "________________"
        content.append({
            "type": "paragraph",
            "text": f"DATED this _____ day of _____________, {execution_date}."
        })
        
        # Signature lines
        content.append({
            "type": "signature_block",
            "grantor": poa_data.grantor,
            "witnesses": poa_data.witnesses
        })
        
        # Witness attestation
        content.append({
            "type": "paragraph",
            "text": self.legal_templates["witness_attestation"]
        })
        
        return content
    
    def _build_personal_care_poa_content(self, poa_data: POADocument) -> List[Dict[str, Any]]:
        """Build content structure for Personal Care POA"""
        content = []
        
        # Title
        content.append({
            "type": "title",
            "text": "POWER OF ATTORNEY FOR PERSONAL CARE"
        })
        
        # Grantor identification and appointment
        attorney_names = ", ".join([a.full_name for a in poa_data.attorneys])
        attorney_addresses = "; ".join([f"{a.address}, {a.city}, {a.province} {a.postal_code}" for a in poa_data.attorneys])
        
        appointment_text = self.legal_templates["personal_care_appointment"].format(
            grantor_name=poa_data.grantor.full_name,
            grantor_address=f"{poa_data.grantor.address}, {poa_data.grantor.city}, {poa_data.grantor.province} {poa_data.grantor.postal_code}",
            attorney_names=attorney_names,
            attorney_addresses=attorney_addresses
        )
        
        content.append({
            "type": "paragraph",
            "text": appointment_text
        })
        
        # Powers granted
        content.append({
            "type": "heading",
            "text": "POWERS GRANTED"
        })
        
        content.append({
            "type": "paragraph",
            "text": self.legal_templates["personal_care_powers"]
        })
        
        # Specific care areas
        if poa_data.personal_care_powers:
            care_areas = []
            if poa_data.personal_care_powers.health_care:
                care_areas.append("Health care decisions, including consent to or refusal of medical treatment")
            if poa_data.personal_care_powers.nutrition:
                care_areas.append("Decisions about nutrition and dietary requirements")
            if poa_data.personal_care_powers.shelter:
                care_areas.append("Decisions about housing and living arrangements")
            if poa_data.personal_care_powers.clothing:
                care_areas.append("Decisions about clothing and personal appearance")
            if poa_data.personal_care_powers.hygiene:
                care_areas.append("Decisions about personal hygiene and cleanliness")
            if poa_data.personal_care_powers.safety:
                care_areas.append("Decisions about personal safety and security")
            if poa_data.personal_care_powers.social_activities:
                care_areas.append("Decisions about social activities and relationships")
            
            if care_areas:
                content.append({
                    "type": "paragraph",
                    "text": "Specifically, my attorney(s) may make decisions about:"
                })
                
                for area in care_areas:
                    content.append({
                        "type": "bullet",
                        "text": area
                    })
        
        # Medical preferences
        if poa_data.personal_care_powers and poa_data.personal_care_powers.medical_preferences:
            content.append({
                "type": "heading",
                "text": "MEDICAL PREFERENCES AND INSTRUCTIONS"
            })
            
            prefs = poa_data.personal_care_powers.medical_preferences
            
            if prefs.get("life_support"):
                content.append({
                    "type": "paragraph",
                    "text": f"Regarding life support: {prefs['life_support']}"
                })
            
            if prefs.get("resuscitation"):
                content.append({
                    "type": "paragraph",
                    "text": f"Regarding resuscitation: {prefs['resuscitation']}"
                })
            
            if prefs.get("organ_donation"):
                content.append({
                    "type": "paragraph",
                    "text": f"Regarding organ donation: {prefs['organ_donation']}"
                })
            
            if prefs.get("religious_considerations"):
                content.append({
                    "type": "paragraph",
                    "text": f"Religious considerations: {prefs['religious_considerations']}"
                })
        
        # Custom instructions
        if poa_data.personal_care_powers and poa_data.personal_care_powers.custom_instructions:
            content.append({
                "type": "heading",
                "text": "ADDITIONAL INSTRUCTIONS"
            })
            
            for instruction in poa_data.personal_care_powers.custom_instructions:
                content.append({
                    "type": "bullet",
                    "text": instruction
                })
        
        # Restrictions and limitations
        content.append({
            "type": "heading",
            "text": "RESTRICTIONS AND LIMITATIONS"
        })
        
        content.append({
            "type": "paragraph",
            "text": "My attorney(s) cannot override any decisions I make while I am capable of making personal care decisions."
        })
        
        content.append({
            "type": "paragraph",
            "text": "My attorney(s) must act in accordance with my prior capable wishes, if known, and otherwise in my best interests."
        })
        
        if poa_data.restrictions and poa_data.restrictions.specific_restrictions:
            for restriction in poa_data.restrictions.specific_restrictions:
                content.append({
                    "type": "bullet",
                    "text": restriction
                })
        
        # Revocation clause
        content.append({
            "type": "paragraph",
            "text": self.legal_templates["revocation_clause"].format(poa_type="Personal Care")
        })
        
        # Capacity statement
        content.append({
            "type": "paragraph",
            "text": self.legal_templates["capacity_statement"]
        })
        
        # Execution section
        content.append({
            "type": "heading",
            "text": "EXECUTION"
        })
        
        execution_date = poa_data.execution_date or "________________"
        content.append({
            "type": "paragraph",
            "text": f"DATED this _____ day of _____________, {execution_date}."
        })
        
        # Signature lines
        content.append({
            "type": "signature_block",
            "grantor": poa_data.grantor,
            "witnesses": poa_data.witnesses
        })
        
        # Witness attestation
        content.append({
            "type": "paragraph",
            "text": self.legal_templates["witness_attestation"]
        })
        
        return content
    
    def _create_pdf_document(self, content: List[Dict[str, Any]], file_path: str, title: str) -> bool:
        """Create PDF document from content structure"""
        try:
            doc = SimpleDocTemplate(file_path, pagesize=letter, topMargin=1*inch, bottomMargin=1*inch)
            story = []
            
            for item in content:
                if item["type"] == "title":
                    story.append(Paragraph(item["text"], self.styles["title"]))
                    story.append(Spacer(1, 12))
                
                elif item["type"] == "heading":
                    story.append(Spacer(1, 12))
                    story.append(Paragraph(item["text"], self.styles["heading"]))
                    story.append(Spacer(1, 6))
                
                elif item["type"] == "paragraph":
                    story.append(Paragraph(item["text"], self.styles["body"]))
                    story.append(Spacer(1, 6))
                
                elif item["type"] == "bullet":
                    bullet_text = f"• {item['text']}"
                    story.append(Paragraph(bullet_text, self.styles["body"]))
                    story.append(Spacer(1, 3))
                
                elif item["type"] == "signature_block":
                    story.append(Spacer(1, 24))
                    
                    # Grantor signature
                    story.append(Paragraph("_" * 50, self.styles["signature"]))
                    story.append(Paragraph(f"{item['grantor'].full_name} (Grantor)", self.styles["signature"]))
                    story.append(Spacer(1, 12))
                    
                    # Witness signatures
                    if item.get("witnesses"):
                        story.append(Paragraph("WITNESSES:", self.styles["heading"]))
                        story.append(Spacer(1, 12))
                        
                        for i, witness in enumerate(item["witnesses"]):
                            story.append(Paragraph("_" * 50, self.styles["signature"]))
                            story.append(Paragraph(f"{witness.full_name}", self.styles["signature"]))
                            story.append(Paragraph(f"{witness.address}, {witness.city}, {witness.province} {witness.postal_code}", self.styles["signature"]))
                            story.append(Spacer(1, 12))
            
            doc.build(story)
            return True
            
        except Exception as e:
            logger.error(f"Error creating PDF document: {e}")
            return False
    
    def _create_word_document(self, content: List[Dict[str, Any]], file_path: str, title: str) -> bool:
        """Create Word document from content structure"""
        try:
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for item in content:
                if item["type"] == "title":
                    # Already added as document title
                    continue
                
                elif item["type"] == "heading":
                    doc.add_heading(item["text"], level=1)
                
                elif item["type"] == "paragraph":
                    para = doc.add_paragraph(item["text"])
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                elif item["type"] == "bullet":
                    doc.add_paragraph(item["text"], style='List Bullet')
                
                elif item["type"] == "signature_block":
                    doc.add_paragraph()  # Add space
                    
                    # Grantor signature
                    doc.add_paragraph("_" * 50)
                    doc.add_paragraph(f"{item['grantor'].full_name} (Grantor)")
                    doc.add_paragraph()
                    
                    # Witness signatures
                    if item.get("witnesses"):
                        doc.add_heading("WITNESSES:", level=2)
                        
                        for witness in item["witnesses"]:
                            doc.add_paragraph("_" * 50)
                            doc.add_paragraph(witness.full_name)
                            doc.add_paragraph(f"{witness.address}, {witness.city}, {witness.province} {witness.postal_code}")
                            doc.add_paragraph()
            
            doc.save(file_path)
            return True
            
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            return False
    
    def get_poa_requirements_info(self, poa_type: str) -> Dict[str, Any]:
        """Get detailed information about POA requirements"""
        if poa_type.lower() == "property":
            return self.ontario_requirements["property_poa"]
        elif poa_type.lower() == "personal_care":
            return self.ontario_requirements["personal_care_poa"]
        else:
            return {
                "property": self.ontario_requirements["property_poa"],
                "personal_care": self.ontario_requirements["personal_care_poa"]
            }
    
    def validate_poa_compliance(self, poa_data: POADocument) -> Dict[str, Any]:
        """Comprehensive POA compliance validation"""
        if poa_data.poa_type == POAType.PROPERTY:
            return self._validate_property_poa_data(poa_data)
        elif poa_data.poa_type == POAType.PERSONAL_CARE:
            return self._validate_personal_care_poa_data(poa_data)
        elif poa_data.poa_type == POAType.COMBINED:
            property_validation = self._validate_property_poa_data(poa_data)
            care_validation = self._validate_personal_care_poa_data(poa_data)
            
            return {
                "valid": property_validation["valid"] and care_validation["valid"],
                "property_validation": property_validation,
                "care_validation": care_validation
            }
        else:
            return {"valid": False, "errors": ["Invalid POA type specified"]}

# Initialize global POA generator
enhanced_poa_generator = EnhancedPOAGenerator()

def get_enhanced_poa_generator() -> EnhancedPOAGenerator:
    """Get the global enhanced POA generator instance"""
    return enhanced_poa_generator

